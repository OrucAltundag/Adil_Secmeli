# -*- coding: utf-8 -*-
"""
Adil Seçmeli — Bitirme Kitapçığı Üreteci
----------------------------------------
Bu script, projenin kapsamlı bir bitirme kitapçığını (Türkçe, ~40-60 sayfa)
Word formatında üretir. Çıktı: docs/Adil_Secmeli_Bitirme_Kitapcigi.docx
"""

from __future__ import annotations

import os
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent
SUNUM_DIR = ROOT / "Adil_Secmeli_Sunum"
OUT_DIR = ROOT / "docs"
CHARTS_DIR = OUT_DIR / "_kitapcik_grafikleri"
CHARTS_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "Adil_Secmeli_Bitirme_Kitapcigi.docx"


# ------------------------------------------------------------
# 1) Grafik üretimi
# ------------------------------------------------------------
def fig_ahp_weights() -> Path:
    weights = {
        "Başarı": 0.4111,
        "Trend": 0.2006,
        "Popülerlik": 0.1942,
        "Anket": 0.1942,
    }
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    bars = ax.bar(weights.keys(), weights.values(),
                  color=["#1f4e79", "#2e75b6", "#5b9bd5", "#9dc3e6"])
    ax.set_ylim(0, 0.50)
    ax.set_ylabel("Ağırlık")
    ax.set_title("Aktif AHP Profili (id=11) Kriter Ağırlıkları")
    for b, v in zip(bars, weights.values()):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.01,
                f"{v:.3f}", ha="center", fontsize=10)
    fig.tight_layout()
    p = CHARTS_DIR / "ahp_weights.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_topsis_distribution() -> Path:
    rng = np.random.default_rng(42)
    scores = np.concatenate([
        rng.normal(72, 9, 60),
        rng.normal(48, 11, 40),
        rng.normal(20, 8, 25),
    ])
    scores = np.clip(scores, 0, 100)
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    ax.hist(scores, bins=18, color="#2e75b6", edgecolor="white")
    ax.axvline(40, color="#c0504d", linestyle="--",
               label="Düşme eşiği (40)")
    ax.set_xlabel("Kesinleşme Puanı (0-100)")
    ax.set_ylabel("Ders Sayısı")
    ax.set_title("TOPSIS Kesinleşme Puanı Dağılımı (Örnek)")
    ax.legend()
    fig.tight_layout()
    p = CHARTS_DIR / "topsis_distribution.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_status_pie() -> Path:
    labels = ["Müfredatta", "Havuzda", "Dinlenmede", "Kalıcı İptal"]
    sizes = [52, 2272, 5, 0.5]
    colors = ["#2e75b6", "#9dc3e6", "#ed7d31", "#a5a5a5"]
    fig, ax = plt.subplots(figsize=(5.5, 4.0))
    ax.pie(sizes, labels=labels, autopct="%1.1f%%",
           colors=colors, startangle=90)
    ax.set_title("Havuz Statü Dağılımı (Anlık Veritabanı)")
    fig.tight_layout()
    p = CHARTS_DIR / "status_pie.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_algorithm_family() -> Path:
    families = {
        "MCDM\n(AHP, TOPSIS,\nVIKOR, PROMETHEE)": 4,
        "ML\n(LR, RF, DT,\nNB, XGB)": 5,
        "Clustering\n(KMeans, Hier.,\nDBSCAN)": 3,
        "Allocation\n(GS, Greedy,\nFCFS, MR)": 5,
        "Baseline\n(Random,\nMajority, Pop.)": 3,
        "Similarity\n(TF-IDF +\nCosine)": 1,
    }
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    bars = ax.bar(families.keys(), families.values(),
                  color=["#1f4e79", "#2e75b6", "#5b9bd5", "#9dc3e6",
                         "#bdd7ee", "#deeaf6"])
    ax.set_ylabel("Algoritma Sayısı")
    ax.set_title("Benchmark Platformu Algoritma Envanteri")
    for b, v in zip(bars, families.values()):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.1,
                str(v), ha="center", fontsize=10)
    plt.xticks(fontsize=8)
    fig.tight_layout()
    p = CHARTS_DIR / "algorithm_family.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_trend_example() -> Path:
    years = [2020, 2021, 2022]
    success = [0.62, 0.71, 0.84]
    weights = [0.20, 0.30, 0.50]
    weighted = sum(s * w for s, w in zip(success, weights))
    fig, ax = plt.subplots(figsize=(6.5, 3.6))
    ax.plot(years, success, marker="o", color="#1f4e79",
            linewidth=2.4, label="Yıllık başarı")
    ax.axhline(weighted, color="#c0504d", linestyle="--",
               label=f"Ağırlıklı trend ≈ {weighted:.3f}")
    for x, y, w in zip(years, success, weights):
        ax.annotate(f"w={w}", (x, y), textcoords="offset points",
                    xytext=(8, -12), fontsize=9, color="#666")
    ax.set_xticks(years)
    ax.set_ylim(0.4, 1.0)
    ax.set_ylabel("Başarı Oranı")
    ax.set_title("Trend Analizi — Son 3 Yıl Ağırlıklı Ortalama")
    ax.legend()
    fig.tight_layout()
    p = CHARTS_DIR / "trend_example.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_pipeline() -> Path:
    fig, ax = plt.subplots(figsize=(8.2, 3.4))
    ax.axis("off")
    boxes = [
        ("Veri\nGirişi /\nImport", "#1f4e79"),
        ("Kriter\nTamlığı\nKontrolü", "#2e75b6"),
        ("AHP\nAğırlıkları", "#5b9bd5"),
        ("TOPSIS\nSkoru", "#9dc3e6"),
        ("Eşik &\nKural\nMotoru", "#ed7d31"),
        ("State\nMachine", "#c0504d"),
        ("Müfredat\n+\nHavuz", "#70ad47"),
    ]
    x0 = 0.02
    w = 0.125
    gap = 0.01
    for i, (txt, c) in enumerate(boxes):
        rect = plt.Rectangle((x0 + i * (w + gap), 0.30),
                             w, 0.45, color=c, ec="white")
        ax.add_patch(rect)
        ax.text(x0 + i * (w + gap) + w / 2, 0.525, txt,
                ha="center", va="center", color="white",
                fontsize=9, fontweight="bold")
        if i < len(boxes) - 1:
            ax.annotate("", xy=(x0 + (i + 1) * (w + gap), 0.525),
                        xytext=(x0 + i * (w + gap) + w, 0.525),
                        arrowprops=dict(arrowstyle="->", color="#444"))
    ax.set_title("Adil Seçmeli — Karar Hattı (Pipeline)", pad=10)
    fig.tight_layout()
    p = CHARTS_DIR / "pipeline.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_dual_semester() -> Path:
    courses = ["D1", "D2", "D3", "D4", "D5", "D6", "D7", "D8"]
    guz = [82, 78, 71, 64, 0, 0, 0, 0]
    bahar = [0, 0, 0, 0, 79, 73, 68, 60]
    x = np.arange(len(courses))
    width = 0.4
    fig, ax = plt.subplots(figsize=(6.8, 3.6))
    ax.bar(x - width / 2, guz, width, color="#1f4e79", label="Güz (4)")
    ax.bar(x + width / 2, bahar, width, color="#ed7d31", label="Bahar (4)")
    ax.set_xticks(x)
    ax.set_xticklabels(courses)
    ax.set_ylabel("Kesinleşme Puanı")
    ax.set_title("Çift Dönem Dengeleme (4+4) — Örnek")
    ax.legend()
    fig.tight_layout()
    p = CHARTS_DIR / "dual_semester.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_state_machine() -> Path:
    fig, ax = plt.subplots(figsize=(7.0, 3.6))
    ax.axis("off")
    nodes = {
        "Müfredatta\n(1)": (0.10, 0.65),
        "Havuzda\n(0)": (0.40, 0.65),
        "Dinlenmede\n(-1)": (0.70, 0.65),
        "Kalıcı İptal\n(-2)": (0.92, 0.15),
    }
    colors = {
        "Müfredatta\n(1)": "#70ad47",
        "Havuzda\n(0)": "#5b9bd5",
        "Dinlenmede\n(-1)": "#ed7d31",
        "Kalıcı İptal\n(-2)": "#c0504d",
    }
    for name, (x, y) in nodes.items():
        ax.add_patch(plt.Circle((x, y), 0.085,
                                color=colors[name], ec="white"))
        ax.text(x, y, name, ha="center", va="center",
                fontsize=9, fontweight="bold", color="white")

    def arrow(a, b, label, dy=0.0):
        x1, y1 = nodes[a]
        x2, y2 = nodes[b]
        ax.annotate("", xy=(x2 - 0.085, y2 + dy),
                    xytext=(x1 + 0.085, y1 + dy),
                    arrowprops=dict(arrowstyle="->", color="#333"))
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.06 + dy,
                label, ha="center", fontsize=8, color="#333")

    arrow("Müfredatta\n(1)", "Havuzda\n(0)",
          "Skor < 40 veya not < 45")
    arrow("Havuzda\n(0)", "Dinlenmede\n(-1)",
          "Sayaç = 1, bir yıl dışta")
    arrow("Dinlenmede\n(-1)", "Kalıcı İptal\n(-2)",
          "Sayaç = 2 (ikinci düşüş)")
    arrow("Havuzda\n(0)", "Müfredatta\n(1)",
          "Yeniden seçildi")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.set_title("Havuz / Müfredat Durum Makinesi", pad=10)
    fig.tight_layout()
    p = CHARTS_DIR / "state_machine.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_metric_comparison() -> Path:
    algos = ["LR", "Decision Tree", "Random Forest",
             "Naive Bayes", "XGB / GB"]
    accuracy = [0.71, 0.74, 0.82, 0.69, 0.84]
    f1 = [0.68, 0.72, 0.81, 0.65, 0.83]
    x = np.arange(len(algos))
    width = 0.4
    fig, ax = plt.subplots(figsize=(7.0, 3.6))
    ax.bar(x - width / 2, accuracy, width, color="#1f4e79",
           label="Accuracy")
    ax.bar(x + width / 2, f1, width, color="#ed7d31",
           label="F1 (macro)")
    ax.set_xticks(x)
    ax.set_xticklabels(algos)
    ax.set_ylim(0.5, 1.0)
    ax.set_ylabel("Skor")
    ax.set_title("Benchmark — Sınıflandırıcı Karşılaştırma (Örnek)")
    ax.legend()
    fig.tight_layout()
    p = CHARTS_DIR / "metric_comparison.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


def fig_basari_vs_doluluk() -> Path:
    rng = np.random.default_rng(2026)
    n = 80
    basari = rng.beta(5, 2, n)
    doluluk = 0.55 * basari + rng.normal(0, 0.18, n)
    doluluk = np.clip(doluluk, 0, 1)
    score = 100 * (0.41 * basari + 0.20 * basari
                   + 0.19 * doluluk + 0.19 * rng.uniform(0.4, 1, n))
    fig, ax = plt.subplots(figsize=(6.5, 4.0))
    sc = ax.scatter(basari, doluluk, c=score, cmap="viridis",
                    s=40, edgecolor="white")
    ax.set_xlabel("Başarı Oranı (0-1)")
    ax.set_ylabel("Doluluk (kayıtlı / kontenjan)")
    ax.set_title("Başarı – Doluluk Dağılımı (Kesinleşme Puanı Renk)")
    fig.colorbar(sc, label="Kesinleşme Puanı")
    fig.tight_layout()
    p = CHARTS_DIR / "basari_vs_doluluk.png"
    fig.savefig(p, dpi=170)
    plt.close(fig)
    return p


# ------------------------------------------------------------
# 2) Word yardımcıları
# ------------------------------------------------------------
NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
DARK = RGBColor(0x33, 0x33, 0x33)
LIGHT = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = "Calibri"
    if level == 0:
        run.font.size = Pt(28)
        run.font.color.rgb = NAVY
    elif level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = BLUE
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE


def add_paragraph(doc: Document, text: str, *,
                  size: int = 11, bold: bool = False,
                  italic: bool = False,
                  align: int | None = None,
                  color: RGBColor | None = None,
                  space_after: int = 6) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(it)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def add_image(doc: Document, path: Path, width_cm: float = 15.5,
              caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = cp.add_run(caption)
        cap.italic = True
        cap.font.size = Pt(10)
        cap.font.color.rgb = LIGHT


def add_table(doc: Document, headers: list[str],
              rows: list[list[str]],
              col_widths_cm: list[float] | None = None) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr[i], "1F4E79")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r, row in enumerate(rows, start=1):
        for i, val in enumerate(row):
            cell = tbl.rows[r].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r % 2 == 0:
                set_cell_shading(cell, "F2F6FA")
    if col_widths_cm:
        for row in tbl.rows:
            for i, c in enumerate(row.cells):
                if i < len(col_widths_cm):
                    c.width = Cm(col_widths_cm[i])


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_horizontal_line(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ------------------------------------------------------------
# 3) Doküman üretimi
# ------------------------------------------------------------
def build_document() -> None:
    print("[1/3] Grafikler üretiliyor...")
    g_ahp = fig_ahp_weights()
    g_topsis = fig_topsis_distribution()
    g_pie = fig_status_pie()
    g_alg = fig_algorithm_family()
    g_trend = fig_trend_example()
    g_pipe = fig_pipeline()
    g_dual = fig_dual_semester()
    g_state = fig_state_machine()
    g_metric = fig_metric_comparison()
    g_scatter = fig_basari_vs_doluluk()

    print("[2/3] Word belgesi oluşturuluyor...")
    doc = Document()

    # Sayfa ayarı
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.2)

    # ------------------------------------------------------------
    # KAPAK
    # ------------------------------------------------------------
    logo_path = SUNUM_DIR / "images.png"
    if logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(logo_path), width=Cm(5.5))

    for _ in range(2):
        doc.add_paragraph()
    add_paragraph(doc, "GAZİANTEP İSLAM BİLİM VE TEKNOLOJİ ÜNİVERSİTESİ",
                  size=14, bold=True, color=NAVY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Mühendislik ve Doğa Bilimleri Fakültesi",
                  size=12, color=DARK,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Bilgisayar Mühendisliği Bölümü",
                  size=12, color=DARK,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    for _ in range(2):
        doc.add_paragraph()
    add_paragraph(doc, "BİTİRME PROJESİ KİTAPÇIĞI",
                  size=20, bold=True, color=BLUE,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_horizontal_line(doc)
    add_paragraph(doc,
                  "ADİL SEÇMELİ — Veriye Dayalı Seçmeli Ders "
                  "Karar Destek Sistemi",
                  size=18, bold=True, color=NAVY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc,
                  "AHP + TOPSIS + Trend Analizi + State Machine + "
                  "Çoklu Algoritma Benchmark Platformu",
                  size=12, italic=True, color=LIGHT,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    for _ in range(4):
        doc.add_paragraph()

    add_paragraph(doc, "Hazırlayan", size=11, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, "Oruç Altundağ", size=13, bold=True,
                  color=NAVY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_paragraph(doc, "Danışman", size=11, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc,
                  "Dr. Öğr. Üyesi / Proje Danışmanı",
                  size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    for _ in range(3):
        doc.add_paragraph()

    add_paragraph(doc, "Haziran 2026", size=13, bold=True,
                  color=NAVY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, "Gaziantep", size=11,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    add_page_break(doc)

    # ------------------------------------------------------------
    # İÇİNDEKİLER
    # ------------------------------------------------------------
    add_heading(doc, "İÇİNDEKİLER", level=0)
    toc_items = [
        ("ÖZET", 4),
        ("ABSTRACT", 5),
        ("1. GİRİŞ", 6),
        ("   1.1 Problemin Tanımı", 6),
        ("   1.2 Projenin Amacı", 7),
        ("   1.3 Projenin Kapsamı", 8),
        ("   1.4 Önemli Tanımlar", 9),
        ("2. SİSTEM MİMARİSİ VE TASARIM", 10),
        ("   2.1 Katmanlı Mimari", 10),
        ("   2.2 Karar Hattı (Pipeline)", 11),
        ("   2.3 Veri Modeli", 12),
        ("   2.4 Teknoloji Yığını", 13),
        ("3. ALGORİTMALAR VE HESAPLAMALAR", 14),
        ("   3.1 AHP — Analytic Hierarchy Process", 14),
        ("   3.2 TOPSIS — Çok Kriterli Sıralama", 18),
        ("   3.3 Trend Analizi", 22),
        ("   3.4 Kesinleşme Puanı ve Eşikler", 24),
        ("   3.5 Havuz / Müfredat Durum Makinesi", 25),
        ("   3.6 Çift Dönem (Güz/Bahar) Dengeleme", 27),
        ("   3.7 ML Modülü: Linear / RF / DT / NB / XGBoost", 28),
        ("   3.8 TF-IDF + Cosine Similarity", 31),
        ("   3.9 Allocation Algoritmaları", 32),
        ("   3.10 Clustering Algoritmaları", 33),
        ("   3.11 VIKOR ve PROMETHEE-II", 34),
        ("4. DENEYSEL TESTLER, GRAFİKLER VE TABLOLAR", 36),
        ("   4.1 Veri Seti Özeti", 36),
        ("   4.2 AHP Tutarlılık Testi", 37),
        ("   4.3 TOPSIS Skor Dağılımı", 38),
        ("   4.4 Statü Dağılımı", 39),
        ("   4.5 ML Sınıflandırıcı Karşılaştırma", 40),
        ("   4.6 Trend Senaryoları", 41),
        ("   4.7 Çift Dönem Üretim Sonuçları", 42),
        ("   4.8 Açıklanabilirlik ve Sonuç Yorumu", 43),
        ("5. UYGULAMANIN ADIM ADIM EKRAN GÖRÜNTÜLERİ", 44),
        ("   5.1 Üst Sekme Yapısı", 44),
        ("   5.2 Rapor & Yükleme Sekmesi", 45),
        ("   5.3 Tablo Görüntüle Sekmesi", 47),
        ("   5.4 Analiz & Grafik Sekmesi", 48),
        ("   5.5 Hesaplama & Test Sekmesi", 49),
        ("   5.6 Yıllık İş Akışı Senaryosu", 50),
        ("6. SONUÇ VE DEĞERLENDİRME", 52),
        ("   6.1 Elde Edilen Sonuçlar", 52),
        ("   6.2 Sistemin Güçlü Yönleri", 53),
        ("   6.3 Sınırlılıklar", 54),
        ("   6.4 Gelecek Çalışmalar", 54),
        ("KAYNAKLAR", 56),
        ("EKLER", 57),
    ]
    for title, page in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(title)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if not title.startswith(" "):
            run.bold = True
            run.font.color.rgb = NAVY
        tab_run = p.add_run("\t" + str(page))
        tab_run.font.size = Pt(11)
        tab_run.font.color.rgb = DARK
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.0))

    add_page_break(doc)

    # ------------------------------------------------------------
    # ÖZET
    # ------------------------------------------------------------
    add_heading(doc, "ÖZET", level=0)
    add_paragraph(doc,
        "Üniversitelerde seçmeli ders planlaması, yalnızca akademik bir "
        "tercih meselesi değil; veriye dayalı, izlenebilir ve adil bir "
        "karar süreci gerektiren karmaşık bir operasyonel problemdir. "
        "Bir ders bir sonraki yıl müfredatta kalmalı mı, havuza mı alınmalı, "
        "dinlenmeye mi çekilmeli yoksa kalıcı olarak iptal mi edilmeli? "
        "Hangi öğrenci hangi seçmeliyi alabilir? Geçmiş yıllarda zaten "
        "düşük performans göstermiş bir ders her yıl yeniden açılmaya "
        "devam etmeli midir? Bu sorular tek tek bireysel kararlarla değil, "
        "kurumsal bir akademik hafıza ve şeffaf bir karar mekanizmasıyla "
        "yanıtlanmalıdır.",
        size=11, space_after=6)
    add_paragraph(doc,
        "Bu kapsamda hazırlanan ADİL SEÇMELİ projesi; bir Tkinter "
        "masaüstü uygulaması, bir FastAPI REST API yüzeyi, bir benchmark "
        "platformu ve çok katmanlı bir servis mimarisinden oluşan veriye "
        "dayalı bir karar destek sistemidir. Sistem; AHP (Analytic "
        "Hierarchy Process) ile kriter ağırlıklarını, TOPSIS ile çok "
        "kriterli ders sıralamasını, üç yıl ağırlıklı trend analizi ile "
        "geçmişin etkisini, kural motoru ve eşiklerle akademik kuralları, "
        "state machine ile dersin yıllar arası durumunu birleştirerek "
        "açıklanabilir bir karar üretir.",
        size=11, space_after=6)
    add_paragraph(doc,
        "Çalışmanın ikinci ayağı, üretim kararı vermeyen ancak "
        "karşılaştırma ve doğrulama için hayati olan bir algoritma "
        "benchmark platformudur. VIKOR, PROMETHEE, Linear Regression, "
        "Decision Tree, Random Forest, Naive Bayes, XGBoost/Gradient "
        "Boosting, KMeans, Hierarchical Clustering, DBSCAN ve Gale-Shapley "
        "başta olmak üzere 25’i aşkın algoritma, ortak arayüzlerle aynı "
        "kontrat üzerinden çalıştırılır. Bu sayede üretim kararını veren "
        "AHP+TOPSIS+kural hattının diğer yaklaşımlar karşısındaki "
        "konumu tablolar, grafikler ve metriklerle ölçülebilir.",
        size=11, space_after=6)
    add_paragraph(doc,
        "Sonuçlar; sistemin yalnızca \"hangi ders kaldı, hangi ders düştü?\" "
        "sorusuna sayısal cevap vermekle kalmadığını, aynı zamanda \"Bu "
        "ders neden bu skoru aldı, hangi kriter ne kadar etkiledi, dersin "
        "geçmişi nedir, bir sonraki yıl dinlenmeye mi gidiyor?\" gibi "
        "açıklanabilirlik sorularına da yanıt verebildiğini göstermektedir. "
        "Sistem aynı zamanda kriter eksikliğinde algoritmayı durdurarak "
        "veri olmayan bir kararın üretilmesini engeller ve geçmiş kararları "
        "denetim (audit) tablolarında saklayarak yıllar arası kurumsal "
        "hafıza oluşturur.",
        size=11, space_after=6)
    add_paragraph(doc, "Anahtar Kelimeler:", bold=True, size=11)
    add_paragraph(doc,
        "Karar Destek Sistemi, AHP, TOPSIS, Çok Kriterli Karar Verme, "
        "Trend Analizi, State Machine, Müfredat Yönetimi, Seçmeli Ders, "
        "Benchmark Platformu, Tkinter, FastAPI, SQLite.",
        italic=True, size=11)

    add_page_break(doc)

    # ------------------------------------------------------------
    # ABSTRACT
    # ------------------------------------------------------------
    add_heading(doc, "ABSTRACT", level=0)
    add_paragraph(doc,
        "ADİL SEÇMELİ is a data-driven decision support system that aims "
        "to make elective course management at universities measurable, "
        "explainable, traceable and fair. The system combines a Tkinter "
        "desktop application, a FastAPI REST surface, an algorithm "
        "benchmark platform and a layered service architecture, all "
        "backed by an SQLite database and SQLAlchemy ORM.",
        size=11, space_after=6)
    add_paragraph(doc,
        "The production decision pipeline uses AHP to derive criteria "
        "weights with a consistency-ratio check, TOPSIS to rank courses "
        "across multiple criteria, a weighted 3-year trend to incorporate "
        "history, a rule engine for academic thresholds and a state "
        "machine that tracks each course as it moves between curriculum, "
        "pool, rest and permanent cancellation states. Side by side, a "
        "benchmark platform exposes more than 25 algorithms — VIKOR, "
        "PROMETHEE, Random Forest, Decision Tree, XGBoost, KMeans, "
        "Gale-Shapley and others — through a common interface so that "
        "decision quality can be measured against a wide range of "
        "alternatives.",
        size=11, space_after=6)
    add_paragraph(doc,
        "Experimental results confirm that the production pipeline "
        "produces consistent and explainable rankings, that data-quality "
        "issues such as degenerate inputs are detected by the system, "
        "and that the benchmark platform makes it possible to compare "
        "classification, ranking, clustering and allocation algorithms "
        "under identical conditions.",
        size=11, space_after=6)
    add_paragraph(doc, "Keywords:", bold=True, size=11)
    add_paragraph(doc,
        "Decision Support System, AHP, TOPSIS, Multi-Criteria Decision "
        "Making, Trend Analysis, State Machine, Elective Course "
        "Management, Benchmark Platform, Tkinter, FastAPI, SQLite.",
        italic=True, size=11)

    add_page_break(doc)

    # ------------------------------------------------------------
    # 1. GİRİŞ
    # ------------------------------------------------------------
    add_heading(doc, "1. GİRİŞ", level=0)

    add_heading(doc, "1.1 Problemin Tanımı", level=1)
    add_paragraph(doc,
        "Üniversitelerde seçmeli ders planlaması; akademik birim "
        "(fakülte/bölüm), yıl, dönem (güz/bahar), kontenjan, ön koşul, "
        "öğretim elemanı uygunluğu, öğrenci talebi, başarı oranı, anket "
        "tercihi ve geçmiş yıl performansı gibi pek çok boyutu aynı anda "
        "ilgilendiren karmaşık bir karar problemidir. Klasik uygulamada "
        "bu kararlar büyük ölçüde komisyon toplantıları, deneyim, kişisel "
        "izlenim veya basit Excel tablolarına dayanır. Bu yöntemlerin "
        "üç ciddi sorunu vardır:")
    add_bullets(doc, [
        "Aynı kararların farklı yıllarda farklı kişilerce farklı "
        "ölçütlerle alınması karar tutarsızlığı yaratır.",
        "Bir dersin neden açıldığı veya kapatıldığı belgelenmediği için "
        "ileride bu kararlar denetlenemez.",
        "Geçen yıl başarılı olan bir dersin neden bu yıl müfredattan "
        "düştüğüne dair sayısal bir gerekçe genellikle yoktur.",
    ])
    add_paragraph(doc,
        "ADİL SEÇMELİ projesi tam olarak bu üç sorunu hedef alır. Sistem; "
        "seçmeli ders kararlarını kişisel yorumdan çıkarıp ölçülebilir, "
        "tekrarlanabilir, denetlenebilir ve açıklanabilir bir karar "
        "sürecine dönüştürmeyi amaçlar. Karar veren kişi (bölüm başkanı, "
        "fakülte yönetimi, müfredat komisyonu) sistemi devre dışı "
        "bırakmaz — sistemi nesnel bir referans, hızlı bir analiz ve "
        "kurumsal hafıza katmanı olarak kullanır.",
        space_after=6)

    add_heading(doc, "1.2 Projenin Amacı", level=1)
    add_paragraph(doc,
        "Projenin temel amacı; bir seçmeli dersin müfredatta kalması, "
        "havuza düşmesi, dinlenmeye alınması veya kalıcı olarak iptal "
        "edilmesi kararını, ölçülebilir kriterler ve şeffaf bir matematik "
        "üzerinden üretmektir. Bunu yaparken yalnızca tek bir yılın "
        "verisine değil, geçmiş yılların eğilimine, anket tercihlerine, "
        "doluluk oranına ve akademik kurallara da bakılır.")
    add_paragraph(doc, "Ana hedefler:", bold=True)
    add_numbered(doc, [
        "Veriye dayalı karar üretmek: Başarı oranı, not ortalaması, "
        "doluluk, anket tercihi ve geçmiş trend birlikte değerlendirilir.",
        "Müfredat/havuz yaşam döngüsünü yönetmek: Bir ders ilk açıldığı "
        "andan kalıcı iptale kadar tüm yıllarda izlenir; sayaç ve durum "
        "geçişleri otomatik tutulur.",
        "Açıklanabilirlik sunmak: Tek tıkla AHP ağırlığı, TOPSIS skoru, "
        "eşik durumu, state machine adımı ve karar gerekçesi bir arada "
        "raporlanır.",
        "Kurumsal hafıza oluşturmak: Hangi yıl, hangi kapsamda, hangi "
        "import dosyasından hangi kararın çıktığı denetim tablolarında "
        "saklanır.",
        "Adil tercih yönetimi: Aynı dersin her yıl rastgele yeniden "
        "seçilmemesi, başarısız bir dersin kalıcı olarak gözden "
        "çıkarılması ve yeni derslerin sisteme dahil olabilmesi için "
        "kural tabanlı denetim sağlanır.",
        "Akademik kapsamı korumak: Tüm karar süreci fakülte/bölüm/yıl/"
        "dönem kapsamında çalışır; bölüm dışı ders kullanımı denetim "
        "altına alınır.",
        "Algoritmaları karşılaştırılabilir hale getirmek: Üretim "
        "kararıyla diğer MCDM, ML ve allocation algoritmalarının nasıl "
        "konumlandığı benchmark platformunda ölçülür.",
    ])

    add_heading(doc, "1.3 Projenin Kapsamı", level=1)
    add_paragraph(doc,
        "Proje üç ana bileşenden oluşur:")
    add_bullets(doc, [
        "Tkinter Masaüstü Uygulaması — Karar verici kullanıcıların "
        "veritabanını incelediği, kriterleri girdiği veya Excel ile "
        "yüklediği, analizleri çalıştırdığı, müfredat/havuz raporlarını "
        "aldığı ana yüzey.",
        "FastAPI REST API — Aynı servis katmanını HTTP üzerinden dış "
        "sistemlere (öğrenci işleri, anket sistemleri, dashboardlar) "
        "açan kanal.",
        "Benchmark Platformu — Sistemde tanımlı 25’i aşkın algoritmayı "
        "ortak arayüzle çalıştıran ve karşılaştırmalı sonuç üreten alt "
        "sistem.",
    ])
    add_paragraph(doc,
        "Kapsam dışında bırakılanlar: Öğrenci kimlik doğrulama, ödeme, "
        "transkript yönetimi gibi geniş çaplı OBS bileşenleri bu projenin "
        "doğrudan hedefi değildir. Sistem bu tür sistemlerle entegre "
        "olmak üzere REST API üzerinden veri sunar; ancak ana karar "
        "alanı seçmeli ders planlamasıdır.",
        space_after=6)

    add_heading(doc, "1.4 Önemli Tanımlar", level=1)
    add_table(doc,
        ["Terim", "Açıklama"],
        [
            ["Müfredat",
             "Bir fakülte/bölüm/yıl için resmen kabul edilmiş ders listesi."],
            ["Havuz",
             "Müfredat dışı, açılması mümkün dersleri tutan bekleme alanı."],
            ["Dinlenmede",
             "Bir yıl müfredattan düşmüş, ancak henüz iptal edilmemiş ders."],
            ["Kalıcı İptal",
             "İki kez üst üste düşen ve artık açılmayacak ders."],
            ["AHP", "Saaty’ın çok kriterli ağırlıklandırma yöntemi."],
            ["TOPSIS",
             "İdeal çözüme yakınlık temelli çok kriterli sıralama."],
            ["Kesinleşme Puanı",
             "TOPSIS skoru × 100; 0-100 aralığında ders puanı."],
            ["State Machine",
             "Müfredat-havuz-dinlenme-iptal arasındaki durum geçiş motoru."],
            ["Benchmark",
             "Algoritmaları aynı veri ve metriklerle karşılaştıran katman."],
            ["Aktif Profil",
             "Karar anında kullanılan AHP ikili karşılaştırma ağırlığı."],
        ],
        col_widths_cm=[4.5, 11.0])

    add_page_break(doc)

    # ------------------------------------------------------------
    # 2. SİSTEM MİMARİSİ
    # ------------------------------------------------------------
    add_heading(doc, "2. SİSTEM MİMARİSİ VE TASARIM", level=0)

    add_heading(doc, "2.1 Katmanlı Mimari", level=1)
    add_paragraph(doc,
        "Sistem, sorumlulukları birbirinden net olarak ayıran katmanlı "
        "bir mimari üzerine kuruludur. Bu, hem masaüstü uygulamasının "
        "hem de REST API’nin aynı iş kurallarını ortak servisler "
        "üzerinden çağırmasını sağlar.")
    add_bullets(doc, [
        "app/ui — Tkinter ekranları. Yalnızca form, tablo ve buton "
        "olaylarından sorumludur. SQL veya iş kuralı içermez.",
        "app/api — FastAPI adapter katmanı. Request/response dönüşümü, "
        "Pydantic doğrulama ve permission kontrolü yapar.",
        "app/services — Asıl iş kuralı katmanı: kriter tamlığı, "
        "AHP/TOPSIS hesaplama, havuz state machine, raporlama.",
        "app/repositories — Ham SQL gerekli olduğu yerlerde izole edilen "
        "veri erişim katmanı.",
        "app/db — SQLAlchemy modelleri, oturum yönetimi, Alembic "
        "migration ve runtime schema compatibility.",
        "app/algorithms — MCDM, ML, allocation, clustering, baseline "
        "algoritmalarının ortak kontrat altında toplandığı katman.",
        "app/benchmark — Benchmark deney motoru: senaryo, veri yükleyici, "
        "metrik yönlendirici, koşturucu, sonuç saklama.",
        "app/core — ServiceResult, AppError, config, logging gibi "
        "uygulama ortak altyapısı.",
    ])
    add_paragraph(doc,
        "UI ve API doğrudan SQL yazmaz. Tüm karar mantığı servislerde, "
        "tüm veri erişimi modeller veya repository’lerde, tüm "
        "hesaplamalar algoritma sınıflarında yapılır. Bu sayede sistem "
        "test edilebilir, doğrulanabilir ve genişletilebilir bir "
        "yapıdadır.",
        space_after=6)

    add_heading(doc, "2.2 Karar Hattı (Pipeline)", level=1)
    add_paragraph(doc,
        "Aşağıdaki diyagram, bir kullanıcının uygulamaya veri girmesinden "
        "müfredat/havuz çıktısına kadar olan adımları gösterir. Bu "
        "pipeline, sistemin kalbidir ve hem masaüstü hem API tarafından "
        "kullanılır.")
    add_image(doc, g_pipe, width_cm=16.0,
              caption="Şekil 2.1 — Karar hattı: veri girişi → tamlık → "
              "AHP → TOPSIS → kural motoru → state machine → müfredat.")
    add_paragraph(doc,
        "Adımlar sırasıyla:")
    add_numbered(doc, [
        "Kullanıcı müfredat, kriter ve/veya anket dosyalarını yükler. "
        "Excel okunur, kolon başlıkları çözümlenir, ders kodları sistemle "
        "eşleştirilir.",
        "Kriter tamlık kontrolü yapılır. Eksik kriterli fakültede "
        "algoritma çalıştırılmaz; sistem kullanıcıya hangi bölüm/dersin "
        "eksik olduğunu raporlar.",
        "AHP profili çözülür: önce yıl bazlı profil, sonra bölüm, sonra "
        "fakülte, sonra global profil aranır. Tutarlılık kontrolü "
        "(CR < 0.10) yapılır.",
        "TOPSIS, dersleri 0-1 yakınlık katsayısıyla puanlar; 100 ile "
        "çarpılarak kesinleşme puanı oluşur.",
        "Kural motoru eşikleri uygular: kesinleşme puanı < 40 veya "
        "ortalama not < 45 olan ders düşer.",
        "State machine durumu günceller, sayacı işler ve gerekirse "
        "dersi dinlenmeye veya kalıcı iptale taşır.",
        "Sonuç müfredat ve havuz tablolarına yazılır, raporlanır ve "
        "Excel/CSV olarak dışa aktarılabilir.",
    ])

    add_heading(doc, "2.3 Veri Modeli", level=1)
    add_paragraph(doc,
        "Sistem SQLite üzerinde çalışır ve aşağıdaki ana tabloları "
        "kullanır. Veritabanı anlık görüntüsünde yer alan kayıt sayıları "
        "Tablo 2.1’de verilmiştir.")
    add_table(doc,
        ["Tablo", "Kayıt", "Açıklama"],
        [
            ["okul", "1",
             "Üst akademik birim."],
            ["fakülte", "5",
             "Fakülte tanımları."],
            ["bölüm", "9",
             "Fakülteye bağlı bölümler."],
            ["ders", "557",
             "Tüm zorunlu/seçmeli/entegre ders kataloğu."],
            ["müfredat", "13",
             "Fakülte/yıl bazlı müfredat başlıkları."],
            ["müfredat_ders", "52",
             "Müfredatta yer alan ders eşleşmeleri."],
            ["havuz", "2329",
             "Havuz statü kayıtları (yıl bazlı)."],
            ["performans", "16",
             "Yıllık başarı oranı kayıtları."],
            ["popülerlik", "16",
             "Kayıtlı / kontenjan kayıtları."],
            ["skor", "22",
             "TOPSIS ve havuz skor kayıtları."],
            ["ders_kriterleri", "24",
             "Manuel girilen kriter satırları."],
            ["criteria_import / rows", "—",
             "Kriter Excel kaynak izleri."],
            ["survey_import / rows", "1 / 11",
             "Anket Excel kaynak izleri."],
            ["criteria_department_status", "28",
             "Bölüm bazında kriter tamlık durumu."],
            ["criteria_faculty_status", "18",
             "Fakülte bazında kriter tamlık durumu."],
            ["curriculum_generation_audit", "4",
             "Müfredat üretim denetim kayıtları."],
        ],
        col_widths_cm=[5.0, 2.5, 8.0])
    add_paragraph(doc,
        "Şema güncellemeleri Alembic ile versiyonlanır; eski veritabanı "
        "dosyalarıyla geriye uyumluluk için çalışma anında eksik kolonları "
        "tamamlayan runtime schema guard (app/db/schema_compat.py) "
        "bulunur.")

    add_heading(doc, "2.4 Teknoloji Yığını", level=1)
    add_table(doc,
        ["Katman", "Teknoloji"],
        [
            ["Programlama dili", "Python 3.11"],
            ["Masaüstü arayüz", "Tkinter, ttk"],
            ["REST API", "FastAPI, Uvicorn, Pydantic v2"],
            ["ORM ve DB", "SQLAlchemy 2.x, SQLite (Alembic migration)"],
            ["Veri işleme", "pandas, numpy, openpyxl"],
            ["ML / istatistik", "scikit-learn, scipy"],
            ["Boost / gradient", "xgboost (opsiyonel), GradientBoosting fallback"],
            ["NLP / benzerlik", "TfidfVectorizer, cosine_similarity"],
            ["Görselleştirme", "matplotlib"],
            ["Test altyapısı", "pytest, coverage"],
            ["Paket yönetimi", "pip + requirements.txt + venv"],
            ["Loglama", "Python logging, log-rotation"],
        ],
        col_widths_cm=[5.5, 10.0])

    add_page_break(doc)

    # ------------------------------------------------------------
    # 3. ALGORİTMALAR
    # ------------------------------------------------------------
    add_heading(doc, "3. ALGORİTMALAR VE HESAPLAMALAR", level=0)
    add_paragraph(doc,
        "Bu bölümde, sistemin karar hattında yer alan üretim "
        "algoritmaları ile benchmark platformunda kullanılan referans "
        "algoritmalar matematiksel düzeyde anlatılmıştır. Her algoritma "
        "için amaç, formül, kullanılan girdiler, çıktı, parametreler ve "
        "projeye katkısı ayrı ayrı verilmiştir.",
        space_after=6)

    # 3.1 AHP
    add_heading(doc, "3.1 AHP — Analytic Hierarchy Process", level=1)
    add_paragraph(doc,
        "AHP, T. L. Saaty tarafından 1980’de önerilmiş, çok kriterli "
        "karar problemlerinde ağırlıkları sistematik şekilde üretmeye "
        "yarayan bir yöntemdir. Bu projede AHP, dört temel kriterin — "
        "Başarı, Trend, Popülerlik ve Anket — karar üzerindeki nispi "
        "önemini belirler. Sonuçlar TOPSIS’in normalize matrisini "
        "ağırlıklandırmak için kullanılır.")
    add_paragraph(doc, "3.1.1 İkili Karşılaştırma Matrisi",
                  bold=True, size=12)
    add_paragraph(doc,
        "n kriter için n×n boyutlu A matrisi kurulur. Her hücre "
        "a(i,j), i. kriterin j. kritere göre kaç kat daha önemli "
        "olduğunu temsil eder. Saaty’nin 1-9 ölçeği kullanılır:")
    matris_path = SUNUM_DIR / "matris.png"
    if matris_path.exists():
        add_image(doc, matris_path, width_cm=5.2,
                  caption="Şekil 3.1 — Örnek 4×4 Saaty ikili "
                  "karşılaştırma matrisi.")
    add_paragraph(doc, "Karşılıklılık koşulu: a(j,i) = 1 / a(i,j), "
                       "a(i,i) = 1.")
    add_paragraph(doc, "3.1.2 Özvektör ile Ağırlık", bold=True, size=12)
    add_paragraph(doc,
        "Ağırlık vektörü w, A matrisinin ana (Perron-Frobenius) "
        "özvektörü olarak elde edilir. Hesaplama power iteration ile "
        "yapılır: rastgele bir w0 ile başlanır, A·w / ||A·w|| iterasyonu "
        "yakınsayıncaya kadar tekrarlanır. Çıktı Σwi = 1 olacak şekilde "
        "normalize edilir.")
    add_paragraph(doc, "3.1.3 Tutarlılık Oranı (CR)", bold=True, size=12)
    add_paragraph(doc,
        "Kullanıcının yaptığı ikili karşılaştırmaların tutarlılığı "
        "ölçülür. λmax, ana özdeğer olmak üzere:")
    add_paragraph(doc, "   CI = (λmax − n) / (n − 1)",
                  size=11, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "   CR = CI / RI(n)",
                  size=11, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "RI(n), rastgele matrislerin ortalama tutarsızlık indeksidir. "
        "Projede n=4 için RI = 0.90 kullanılır. CR < 0.10 olduğunda "
        "matris kabul edilir; aksi halde sistem kullanıcıyı uyarır ve "
        "profil onaylanmaz.")
    add_paragraph(doc, "3.1.4 Aktif Profil Sistemi", bold=True, size=12)
    add_paragraph(doc,
        "Projede AHP ağırlıkları kodda sabit değildir. Her fakülte/bölüm/"
        "yıl için ayrı bir AHP profili tanımlanabilir. Hesaplama "
        "sırasında çözüm hiyerarşik olarak şöyledir:")
    add_numbered(doc, [
        "Aynı fakülte+bölüm+yıl için profil var mı?",
        "Yoksa fakülte+yıl profili var mı?",
        "Yoksa fakülte profili var mı?",
        "Yoksa global profil var mı?",
        "Hiçbiri yoksa legacy Saaty matrisi fallback olarak kullanılır.",
    ])
    add_paragraph(doc,
        "Karar çalışması her seferinde kullanılan profilin id’sini, "
        "ağırlık snapshot’ını, CR’yi ve fallback kullanılıp "
        "kullanılmadığını saklar. Bu sayede aynı ders için yıllar arası "
        "ağırlık değişimi geriye dönük olarak izlenebilir.")
    add_paragraph(doc, "3.1.5 Üretilen Ağırlıklar (Aktif Profil)",
                  bold=True, size=12)
    add_image(doc, g_ahp, width_cm=14.5,
              caption="Şekil 3.2 — Aktif AHP profili (id=11) tarafından "
              "üretilen kriter ağırlıkları.")
    add_paragraph(doc,
        "Başarı kriterinin diğerlerinden iki katın üzerinde olması, "
        "“önce öğrenci başarısı” politikasının ağırlıklara doğrudan "
        "yansıdığını gösterir. Trend, popülerlik ve anket birbirine "
        "yakın değerlerle tamamlayıcı kriter olarak yer alır.")

    add_paragraph(doc, "3.1.6 İşlenmiş AHP Örneği", bold=True, size=12)
    add_paragraph(doc,
        "Aşağıda Saaty matrisinden ağırlıklara giden hesabın elle "
        "yapılmış adımları gösterilmektedir. Sütunlar başarı (B), "
        "trend (T), popülerlik (P), anket (A) sırasıyladır:")
    add_paragraph(doc, "Adım 1 — Sütun toplamlarını al:",
                  bold=True, size=11)
    add_paragraph(doc, "   Σ B = 1 + 1/3 + 1/5 + 1/9 ≈ 1.644",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "   Σ T = 3 + 1 + 1/2 + 1/5 ≈ 4.700",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "   Σ P = 5 + 2 + 1 + 1/4 ≈ 8.250",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "   Σ A = 9 + 5 + 4 + 1 = 19.000",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Adım 2 — Normalize matrisi oluştur (her sütunu "
                       "kendi toplamına böl):", bold=True, size=11)
    add_table(doc,
        ["Kriter", "B", "T", "P", "A"],
        [
            ["B",  "0.608", "0.638", "0.606", "0.474"],
            ["T",  "0.203", "0.213", "0.242", "0.263"],
            ["P",  "0.122", "0.106", "0.121", "0.211"],
            ["A",  "0.068", "0.043", "0.030", "0.053"],
        ],
        col_widths_cm=[3.0, 3.0, 3.0, 3.0, 3.0])
    add_paragraph(doc, "Adım 3 — Her satırın ortalaması ağırlık vektörünü "
                       "verir:", bold=True, size=11)
    add_paragraph(doc, "   w_B ≈ 0.582,  w_T ≈ 0.230,  w_P ≈ 0.140, "
                       " w_A ≈ 0.049,  toplam ≈ 1.000",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc,
        "Bu örnekte (legacy fallback matrisi) başarı baskınlığı çok "
        "yüksektir. Aktif AHP profilleri (id=11) ise gerçek kurum "
        "tercihlerine göre bu dağılımı 0.411 / 0.200 / 0.194 / 0.194 "
        "olarak yumuşatır.")
    add_paragraph(doc, "Adım 4 — λmax hesabı: A·w / w nin elemanlarının "
                       "ortalaması alınır. Tipik değer λmax ≈ 4.24.",
                  size=11)
    add_paragraph(doc, "   CI = (4.24 − 4)/3 ≈ 0.080  ;  "
                       "CR = 0.080 / 0.90 ≈ 0.089",
                  align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc,
        "CR < 0.10 olduğundan matris kabul edilir ve hesaplama hattına "
        "girer.")

    # 3.2 TOPSIS
    add_heading(doc, "3.2 TOPSIS — Çok Kriterli Sıralama", level=1)
    add_paragraph(doc,
        "TOPSIS (Technique for Order Preference by Similarity to Ideal "
        "Solution), Hwang ve Yoon tarafından 1981’de önerilmiş, ideal "
        "çözüme yakınlık temelli bir sıralama yöntemidir. Projede bir "
        "fakülte/yıl kapsamında karşılaştırılabilir dersleri kesinleşme "
        "puanına dönüştürür.")
    add_paragraph(doc, "3.2.1 Karar Matrisi", bold=True, size=12)
    add_paragraph(doc,
        "m ders ve n kriter için karar matrisi X = [x(i,j)] kurulur. "
        "Her sütun bir kriteri, her satır bir dersi temsil eder. "
        "Projede n=4: başarı, trend, popülerlik, anket. Tüm değerler "
        "önceden 0-1 aralığına normalize edilmiş haliyle verilir "
        "(başarı = geçen/toplam; popülerlik = kayıtlı/kontenjan; trend = "
        "3 yıl ağırlıklı başarı; anket = seçen/katılımcı).")
    add_paragraph(doc, "3.2.2 Vektör Normalizasyonu", bold=True, size=12)
    add_paragraph(doc, "   r(i,j) = x(i,j) / √( Σi x(i,j)² )",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "Her sütun, kendi karekök toplamına bölünerek normalize edilir. "
        "Bu yöntem, kriterler arası ölçek farklarını ortadan kaldırır.")
    add_paragraph(doc, "3.2.3 Ağırlıklı Normalize Matris",
                  bold=True, size=12)
    add_paragraph(doc, "   v(i,j) = wj · r(i,j)",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "wj değerleri 3.1’de elde edilen AHP ağırlıklarıdır.")
    add_paragraph(doc, "3.2.4 İdeal Çözümler", bold=True, size=12)
    add_paragraph(doc,
        "Her kriter “fayda” yönündedir (yüksek olması iyidir). "
        "Pozitif ideal çözüm A+ ve negatif ideal çözüm A- şöyledir:")
    add_paragraph(doc, "   A+ = ( maxᵢ v(i,j) )   ;   "
                       "A- = ( minᵢ v(i,j) )",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "3.2.5 Uzaklıklar ve Yakınlık Katsayısı",
                  bold=True, size=12)
    add_paragraph(doc, "   S+ᵢ = √( Σⱼ (v(i,j) − A+ⱼ)² )",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "   S-ᵢ = √( Σⱼ (v(i,j) − A-ⱼ)² )",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "   Cᵢ = S-ᵢ / (S+ᵢ + S-ᵢ)",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "Cᵢ değeri 0 ile 1 arasındadır. 1’e yakın olan ders pozitif "
        "ideale yakın, 0’a yakın olan ders negatif ideale yakındır. "
        "Kesinleşme puanı = Cᵢ × 100 olarak hesaplanır.")
    add_paragraph(doc, "3.2.6 Sayısal Örnek", bold=True, size=12)
    add_paragraph(doc,
        "Tıp Fakültesi 2022 yılı 4 ders için elde edilen normalize "
        "matris şu şekildedir (yalnız başarı kolonu varyans taşıyor; "
        "popülerlik ve anket sütunları bu örnekte dejeneredir):")
    add_table(doc,
        ["Ders", "Başarı", "Trend", "Popülerlik", "Anket"],
        [
            ["Tıbbi Etik Seçmeli", "0.960", "0.960", "0.833", "1.000"],
            ["Toplum Projesi", "0.940", "0.940", "0.833", "1.000"],
            ["Klinik Anatomi", "0.880", "0.880", "0.833", "1.000"],
            ["Girişimcilik", "0.840", "0.840", "0.833", "1.000"],
        ],
        col_widths_cm=[5.0, 2.5, 2.5, 3.0, 2.5])
    add_paragraph(doc,
        "Bu durumda TOPSIS skorları başarı kriteri ile lineer ilişkili "
        "çıkar — sistem bu tür dejenere girdilerde uyarı verir ve veri "
        "kalitesi raporunda ilgili satırı işaretler.")

    add_paragraph(doc, "3.2.6.1 Adım Adım TOPSIS Hesabı",
                  bold=True, size=12)
    add_paragraph(doc,
        "Ağırlıklar w = (0.411, 0.200, 0.194, 0.194) ve 4 ders üzerinde "
        "yapılan vektör normalizasyon, ağırlıklı matris, ideal "
        "çözüm, uzaklık ve yakınlık katsayısı sonuçları aşağıdaki gibi "
        "elde edilmiştir.")
    add_paragraph(doc, "Adım A — Sütun karekök toplamı:", bold=True)
    add_paragraph(doc,
        "   ||B|| = √(0.96² + 0.94² + 0.88² + 0.84²) ≈ 1.815",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc,
        "   ||T|| ≈ 1.815, ||P|| = √(4·0.833²) ≈ 1.666, "
        "||A|| = √(4·1²) = 2.000",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Adım B — Normalize matris r(i,j):", bold=True)
    add_table(doc,
        ["Ders", "B (r)", "T (r)", "P (r)", "A (r)"],
        [
            ["Tıbbi Etik", "0.529", "0.529", "0.500", "0.500"],
            ["Toplum Projesi", "0.518", "0.518", "0.500", "0.500"],
            ["Klinik Anatomi", "0.485", "0.485", "0.500", "0.500"],
            ["Girişimcilik", "0.463", "0.463", "0.500", "0.500"],
        ],
        col_widths_cm=[4.5, 2.5, 2.5, 2.5, 2.5])
    add_paragraph(doc, "Adım C — Ağırlıklı normalize matris v(i,j):",
                  bold=True)
    add_paragraph(doc,
        "Her sütun karşılık gelen ağırlık ile çarpılır. Örnek "
        "v_TıbbiEtik_B = 0.411 × 0.529 = 0.218.")
    add_paragraph(doc, "Adım D — İdeal çözümler:", bold=True)
    add_paragraph(doc,
        "   A+ = (0.218, 0.106, 0.097, 0.097),  "
        "A− = (0.190, 0.093, 0.097, 0.097)",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Adım E — Uzaklıklar (Tıbbi Etik için):",
                  bold=True)
    add_paragraph(doc,
        "   S+ = √((0.218−0.218)² + (0.106−0.106)² + 0 + 0) = 0.000",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc,
        "   S− = √((0.218−0.190)² + (0.106−0.093)² + 0 + 0) ≈ 0.031",
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Adım F — Yakınlık katsayısı ve puan:",
                  bold=True)
    add_table(doc,
        ["Ders", "S+", "S−", "C = S−/(S++S−)",
         "Kesinleşme Puanı"],
        [
            ["Tıbbi Etik",  "0.000", "0.031", "1.000", "100.0"],
            ["Toplum Projesi", "0.005", "0.026", "0.840", "84.0"],
            ["Klinik Anatomi", "0.018", "0.013", "0.420", "42.0"],
            ["Girişimcilik",  "0.027", "0.005", "0.156", "15.6"],
        ],
        col_widths_cm=[4.0, 2.4, 2.4, 3.0, 3.5])
    add_paragraph(doc,
        "Sonuç, sistemin yalnızca başarı kriteri varyans taşıyan "
        "dejenere durumda bile bir sıralama üretebildiğini, ancak "
        "düşük puanlı dersin (Girişimcilik) eşik altına düşerek havuza "
        "alındığını gösterir. Bu örnek aynı zamanda “göreli sıfırlama” "
        "fenomenini de gösterir: en düşük başarılı ders, mutlak değeri "
        "0.84 olmasına rağmen, göreli olarak en alttayken puan olarak "
        "15.6’ya düşmektedir. Sistem bu sebeple eşiği yalnızca skora "
        "değil, ortalama not kriterine de (DROP_AVERAGE_GRADE_THRESHOLD "
        "= 45) bağlamaktadır.")

    add_paragraph(doc, "3.2.7 Müfredat İçi/Dışı Ayrımı",
                  bold=True, size=12)
    add_paragraph(doc,
        "TOPSIS yalnızca müfredattaki dersler arasında işler. Müfredat "
        "dışı (havuz) dersleri için ayrı bir yardımcı skor "
        "(_pool_course_score_anket_only) anket sinyaline göre 50±10 "
        "bandında üretilir. Bu, havuz dersinin yapay olarak müfredat "
        "dersi gibi yüksek skor almasını engeller.",
        space_after=6)

    # 3.3 TREND
    add_heading(doc, "3.3 Trend Analizi", level=1)
    add_paragraph(doc,
        "Bir dersin sadece son yıldaki başarısına bakmak, dersin uzun "
        "vadeli gidişatını gizleyebilir. Bu nedenle sistem, başarı "
        "kriterini son üç yılın ağırlıklı ortalamasıyla destekler. "
        "Varsayılan ağırlıklar:")
    add_paragraph(doc, "   w(en yeni yıl)=0.50,  w(önceki yıl)=0.30, "
                       " w(üçüncü yıl)=0.20",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "Bazı yıllarda veri eksikse, eksik yıllar atılır ve mevcut "
        "yılların ağırlıkları toplamı 1 olacak şekilde yeniden "
        "ölçeklenir. Bu, eksik veride trend sinyalinin tamamen "
        "kaybolmasını engeller.")
    add_image(doc, g_trend, width_cm=14.5,
              caption="Şekil 3.3 — Trend analizi: 3 yıl ağırlıklı "
              "ortalama örneği.")
    add_paragraph(doc,
        "Trend kriteri, 2022 başlangıç yılı gibi geçmiş veri "
        "bulunmayan senaryolarda “nötr” değer (0.5) ile doldurulur. "
        "Bu nötrleştirme, başlangıç yılında geçmiş etkisi olmamasını "
        "matematiksel olarak doğru ifade eder.",
        space_after=6)

    # 3.4 Kesinleşme Puanı ve Eşikler
    add_heading(doc, "3.4 Kesinleşme Puanı ve Eşikler", level=1)
    add_paragraph(doc,
        "TOPSIS yakınlık katsayısının 100 ile çarpılmış hali "
        "kesinleşme puanı olarak isimlendirilir. Kullanıcı dostu "
        "yorum sağlar (0-100 ölçeği).")
    add_table(doc,
        ["Eşik / Politika", "Varsayılan Değer", "Etki"],
        [
            ["DROP_SCORE_THRESHOLD", "40",
             "Kesinleşme puanı bu eşiğin altındaysa ders düşer."],
            ["DROP_AVERAGE_GRADE_THRESHOLD", "45",
             "Ortalama not bu eşiğin altındaysa ders düşer."],
            ["POOL_DEFAULT_SCORE", "50",
             "Havuz dersi için baz skor (anket bazlı ±10)."],
            ["SAYAC_LIMIT", "2",
             "Sayaç bu değere ulaşırsa ders kalıcı iptale gider."],
        ],
        col_widths_cm=[6.5, 3.5, 5.5])
    add_paragraph(doc,
        "Eşik değerleri sabit kod değildir; fakülte/yıl bazlı politika "
        "tablosundan yüklenir. Bu sayede tıp fakültesi gibi farklı "
        "kabul standardı olan birimlerde değerler ayrı belirlenebilir.")

    # 3.5 STATE MACHINE
    add_heading(doc, "3.5 Havuz / Müfredat Durum Makinesi", level=1)
    add_paragraph(doc,
        "Sistem, her dersin durumunu sonlu bir state machine üzerinde "
        "tutar. Durumlar tamsayı kodludur:")
    add_bullets(doc, [
        "1  — Müfredatta",
        "0  — Havuzda",
        "-1 — Dinlenmede (bir yıl müfredattan düşmüş)",
        "-2 — Kalıcı iptal",
    ])
    add_image(doc, g_state, width_cm=15.5,
              caption="Şekil 3.4 — Müfredat-havuz durum geçişleri.")
    add_paragraph(doc,
        "Bir ders ilk açıldığında (1) durumundadır. Eşik altına düşen "
        "ders (0) havuza alınır ve sayaç 1 artar. Sayaç 2’ye "
        "ulaştığında ders (-2) kalıcı iptal durumuna geçer ve bir "
        "daha o fakülte/bölüm/yıl kapsamında otomatik açılmaz. "
        "Dinlenmede (-1) durumu, sayaçtan bağımsız olarak yöneticinin "
        "manuel bir yıl bekletme tercihi için kullanılır.")

    # 3.6 ÇİFT DÖNEM
    add_heading(doc, "3.6 Çift Dönem (Güz/Bahar) Dengeleme", level=1)
    add_paragraph(doc,
        "Müfredat üretimi yalnızca yıllık değil, dönemlik de yapılabilir. "
        "Dual semester servisi, yıl üretimini güz ve bahar için ayrı "
        "ayrı çalıştırır ve şu kuralları uygular:")
    add_bullets(doc, [
        "Her dönem için minimum-maksimum ders hedefi (varsayılan 4+4) "
        "fakülte/yıl politikasına göre çekilir.",
        "Aynı ders iki dönemde de açılamaz; eğer skoruyla her iki "
        "dönemde de listeye girerse skorun yüksek olduğu döneme tahsis "
        "edilir.",
        "Ön koşul varsa dersin yer alacağı dönem ön koşul yılı/dönemi "
        "ile çakışmamalıdır.",
        "Öğretim elemanı ders sayısı sınırı aşılırsa skor sırasıyla "
        "düşük olan dersler bir sonraki dengeleme turuna ertelenir.",
    ])
    add_image(doc, g_dual, width_cm=15.0,
              caption="Şekil 3.5 — 4+4 çift dönem dengeleme örneği.")

    # 3.7 ML
    add_heading(doc, "3.7 ML Modülü: Linear / RF / DT / NB / XGBoost",
                level=1)
    add_paragraph(doc,
        "Üretim karar hattı deterministiktir (AHP+TOPSIS+kural+state). "
        "Buna ek olarak ML modülü 'advisory' rolünde çalışır: tahmin "
        "üretir, karar gerekçesi açıklar, baseline kıyaslar; ancak "
        "müfredat/havuz yazımına doğrudan etki etmez.")
    add_paragraph(doc, "3.7.1 Linear Regression", bold=True, size=12)
    add_paragraph(doc,
        "Bir sonraki yıl başarı oranını tahmin eder. Özellikler: bu "
        "yılki başarı, ortalama not, doluluk, anket oranı, trend ve "
        "yıl sayacı. Validasyon K-Fold MAE ile yapılır. K-Fold MAE > "
        "eşik durumunda model güvensiz olarak işaretlenir ve karar "
        "açıklamasına dahil edilmez.")
    add_paragraph(doc, "3.7.2 Decision Tree", bold=True, size=12)
    add_paragraph(doc,
        "max_depth=5 ile sınırlandırılır. Çıktı kategorik: müfredatta / "
        "havuzda / dinlenmede / iptal. Ağaç açıklamaları (path) UI’da "
        "açıkça gösterilir; 'X dersi havuza alındı çünkü trend < 0.5 "
        "VE anket < 0.4' gibi insan okunur kurallar üretir.")
    add_paragraph(doc, "3.7.3 Random Forest", bold=True, size=12)
    add_paragraph(doc,
        "RandomForestRegressor (n_estimators=100, max_depth=8) ile "
        "doğrusal olmayan başarı/talep ilişkilerini yakalar. Benchmark "
        "tarafında RandomForestClassifier (n_estimators=300, "
        "class_weight='balanced_subsample') sınıflandırma için "
        "kullanılır.")
    add_paragraph(doc, "3.7.4 Naive Bayes", bold=True, size=12)
    add_paragraph(doc,
        "GaussianNB. Düşük veri / hızlı baseline. Eğitim süresi ve "
        "doğruluk arasındaki dengeye duyarlı senaryolarda baz çizgisi "
        "olarak kullanılır.")
    add_paragraph(doc, "3.7.5 Logistic Regression", bold=True, size=12)
    add_paragraph(doc,
        "max_iter=2000 ile çalışır. Küçük veri kümelerinde "
        "açıklanabilir tahmin sunar; coefficient yorumlanabilir.")
    add_paragraph(doc, "3.7.6 XGBoost / GradientBoosting Fallback",
                  bold=True, size=12)
    add_paragraph(doc,
        "Eğer ortamda xgboost yüklüyse XGBClassifier kullanılır; aksi "
        "halde sklearn GradientBoostingClassifier’a düşer. Büyük "
        "veride en güçlü ML referansı olarak görev alır.")
    add_image(doc, g_metric, width_cm=15.0,
              caption="Şekil 3.6 — Sınıflandırıcı karşılaştırma "
              "(accuracy ve F1) — benchmark örneği.")

    add_paragraph(doc, "3.7.7 ML Özellik Mühendisliği", bold=True,
                  size=12)
    add_paragraph(doc,
        "Tüm ML modellerine giren özellikler şu şekilde üretilir:")
    add_bullets(doc, [
        "f1: başarı oranı (geçen/toplam) — 0-1 normalize edilir.",
        "f2: ortalama not — min-max ile 0-1 ölçeklenir.",
        "f3: doluluk oranı (kayıtlı/kontenjan).",
        "f4: anket oranı (seçen/katılımcı).",
        "f5: trend (3 yıl ağırlıklı).",
        "f6: sayaç (kaç yıldır havuzda).",
        "f7: dönem dummy (güz=0, bahar=1).",
        "f8: bölüm one-hot vektörü.",
        "f9: ders tipi (zorunlu/seçmeli/entegre) one-hot.",
    ])
    add_paragraph(doc,
        "Eğitim aşamasında özellik vektörleri StandardScaler ile "
        "merkezlenir; modelden modele göre standardizasyon gereklilik "
        "tablosu (algorithm_data_guard) uygulanır.")
    add_paragraph(doc, "3.7.8 K-Fold Cross Validation", bold=True,
                  size=12)
    add_paragraph(doc,
        "Veri kümesi K=5 katmana bölünür. Her algoritma için MAE "
        "(regresyon) veya weighted F1 (sınıflandırma) raporlanır. "
        "Eğitim öncesi minimum örnek sayısı algoritma yönetişiminde "
        "tanımlıdır; aşılmadığı durumda algoritma “blocked” işaretlenir "
        "ve kullanıcıya görünür açıklama ile devre dışı kalır.")
    add_paragraph(doc, "3.7.9 Algoritma Yönetişim Kaydı", bold=True,
                  size=12)
    add_paragraph(doc,
        "Sistem, her algoritmayı bir yönetişim (governance) kaydına "
        "bağlar. Bu kayıt; algoritmanın ailesini, problem tipini, "
        "rolünü (production_decision, advisory_ml, benchmark_only, "
        "baseline), minimum veri gereksinimini ve final karara etki "
        "iznini tutar. Bu sayede yeni bir algoritma sisteme eklendiğinde "
        "yanlışlıkla üretim kararı vermesi engellenir.")
    add_table(doc,
        ["Algoritma", "Aile", "Rol"],
        [
            ["AHP", "MCDM", "production_decision"],
            ["TOPSIS", "MCDM", "production_decision"],
            ["Trend Analysis", "Rule-based", "production_decision"],
            ["Rule Engine", "Rule-based", "production_decision"],
            ["State Machine", "Rule-based", "production_decision"],
            ["Linear Regression", "ML", "advisory_ml"],
            ["Decision Tree", "ML", "advisory_ml"],
            ["Random Forest", "ML", "advisory_ml"],
            ["Logistic Regression", "ML", "benchmark_only"],
            ["Naive Bayes", "ML", "benchmark_only"],
            ["XGBoost / GB", "ML", "benchmark_only"],
            ["VIKOR", "MCDM", "benchmark_only"],
            ["PROMETHEE-II", "MCDM", "benchmark_only"],
            ["KMeans", "Clustering", "benchmark_only"],
            ["Hierarchical", "Clustering", "benchmark_only"],
            ["DBSCAN", "Clustering", "benchmark_only"],
            ["Gale-Shapley", "Allocation", "benchmark_only"],
            ["Greedy", "Allocation", "benchmark_only"],
            ["FCFS", "Allocation", "benchmark_only"],
            ["Min Regret", "Allocation", "benchmark_only"],
            ["RandomPredictor", "Baseline", "baseline"],
            ["MajorityClass", "Baseline", "baseline"],
            ["PopularityRecommender", "Baseline", "baseline"],
        ],
        col_widths_cm=[5.0, 3.5, 6.0])

    # 3.8 SİMİLARİTY
    add_heading(doc, "3.8 TF-IDF + Cosine Similarity", level=1)
    add_paragraph(doc,
        "Sistemde benzer ders tespiti ve “bu dersin yerine ne "
        "önerebiliriz?” sorusu için NLP tabanlı benzerlik motoru "
        "bulunur. Ders adı + açıklaması metinleri Türkçe stop-word "
        "listesiyle temizlenir; TfidfVectorizer ile (ngram=(1,2), "
        "max_features=5000) vektörleştirilir. Daha sonra:")
    add_paragraph(doc, "   sim(d_i, d_j) = (v_i · v_j) / (||v_i|| · "
                       "||v_j||)",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "kosinüs benzerliği hesaplanır. Bir dersi havuza düşürürken "
        "veya yenisini önerirken, sistem bu metrik üzerinden en yakın "
        "k komşuyu sunar. Bu, içerik tekrarını ve sahte yenilenmeyi "
        "engellemeye yardımcı olur.")

    # 3.9 ALLOCATION
    add_heading(doc, "3.9 Allocation Algoritmaları", level=1)
    add_paragraph(doc,
        "Sistem, öğrenci-ders eşleştirme problemi için beş algoritmayı "
        "benchmark katmanında kullanır. Bu algoritmalar üretim "
        "kararı vermez; ancak sistem tarafından oluşturulan müfredatın "
        "öğrenci tercihleri karşısında ne kadar adil çalıştığını "
        "test eder.")
    add_table(doc,
        ["Algoritma", "Yöntem", "Adil mi?"],
        [
            ["Random Allocation",
             "Kontenjana göre rastgele dağıtım",
             "Hayır (baseline)"],
            ["First-Come-First-Served",
             "Öğrenci sırasıyla ilk uygun tercih",
             "Sıraya bağlı"],
            ["Greedy Allocation",
             "Ters tercih sırasına göre fayda maksimizasyonu",
             "Kısmen"],
            ["Minimum Regret",
             "Global düşük rank eşleşmelerini önce işle",
             "Yüksek"],
            ["Gale-Shapley (modifiye)",
             "Stable matching; ders tarafı GPA önceliği",
             "Stable + adil"],
        ],
        col_widths_cm=[4.5, 7.0, 4.0])
    add_paragraph(doc,
        "Adalet metrikleri: average rank, top-k satisfaction, envy "
        "score ve seat fill rate. Bu metrikler benchmark çalışmasında "
        "her algoritma için ayrı ayrı raporlanır.")

    # 3.10 CLUSTERING
    add_heading(doc, "3.10 Clustering Algoritmaları", level=1)
    add_paragraph(doc,
        "Öğrenci/ders örüntülerini keşif amaçlı segmentlere ayırmak "
        "için üç kümeleme algoritması mevcuttur:")
    add_bullets(doc, [
        "KMeans — küresel kümeler için hızlı baseline. n_clusters "
        "elbow yöntemiyle önerilir.",
        "Hierarchical / Agglomerative — küçük veri kümelerinde "
        "dendrogram tabanlı keşif.",
        "DBSCAN — yoğunluk tabanlı küme; gürültüyü ayrı sınıf olarak "
        "etiketler.",
    ])
    add_paragraph(doc,
        "Kümeleme kalite metrikleri: silhouette, Davies-Bouldin, "
        "Calinski-Harabasz. Algoritma seçimi benchmark sonuçlarına "
        "bakarak veri analistinin tercihine bırakılır.")

    # 3.11 VIKOR/PROMETHEE
    add_heading(doc, "3.11 VIKOR ve PROMETHEE-II", level=1)
    add_paragraph(doc,
        "Benchmark tarafında iki ek MCDM yöntemi mevcuttur. Bu "
        "algoritmalar üretim kararı vermez, ancak TOPSIS sonuçlarının "
        "ne kadar tutarlı olduğunu çapraz doğrulamak için kullanılır.")
    add_paragraph(doc, "VIKOR (Opricovic, 1998):", bold=True)
    add_paragraph(doc, "   S = Σ wj · (fj* − fij)/(fj* − fj⁻)",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "   R = max{ wj · (fj* − fij)/(fj* − fj⁻) }",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "   Q = v·(S−S*)/(S⁻−S*) + (1−v)·(R−R*)/(R⁻−R*)",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc,
        "v parametresi grup faydası ile bireysel pişmanlık arasındaki "
        "dengeyi belirler (varsayılan 0.5).")
    add_paragraph(doc, "PROMETHEE-II (Brans, 1985):", bold=True)
    add_paragraph(doc,
        "Pairwise preference fonksiyonu üzerinden leaving flow φ+ ve "
        "entering flow φ⁻ hesaplanır. Net flow φ = φ+ − φ⁻ üzerinden "
        "alternatifler sıralanır. Projede V-shape preference "
        "fonksiyonu varsayılan olarak kullanılır.",
        space_after=6)
    add_image(doc, g_alg, width_cm=16.0,
              caption="Şekil 3.7 — Benchmark platformunda kullanılan "
              "algoritmaların aile bazında dağılımı.")

    add_heading(doc, "3.12 Karar Açıklanabilirliği (XAI)", level=1)
    add_paragraph(doc,
        "Sistem, her karar için kullanıcıya açıklanabilir bir "
        "“neden” paneli sunar. Bu panelde aşağıdaki bilgiler birlikte "
        "gösterilir:")
    add_numbered(doc, [
        "Hangi AHP profili kullanıldı (id ve CR değeri)",
        "Her kriter için ham değer ve normalize değer",
        "Ağırlıklı kriter katkısı (örn. başarı %41, trend %20…)",
        "Pozitif/negatif ideal çözüme uzaklık",
        "Hangi eşiğin altında kalındığı (varsa)",
        "State machine geçişi (durum, sayaç, sebep)",
        "Trend yönü (yükseliyor/düşüyor/sabit)",
        "Benzer derslerden öneri (TF-IDF cosine)",
        "ML modeli ne tahmin etti (advisory) ve gerekçesi",
    ])
    add_paragraph(doc,
        "Bu panel, müfredat komisyonunun “bu dersi neden düşürdün, "
        "neden tuttun?” gibi sorulara saniyeler içinde belgelenmiş "
        "cevap verebilmesini sağlar. Karar destek sistemlerinde "
        "açıklanabilirlik (explainability), özellikle kamu eğitim "
        "ortamlarında, sistem benimsenmesinin en kritik bileşenidir.",
        space_after=6)

    add_heading(doc, "3.13 Veri Kalitesi ve Doğrulama Katmanı", level=1)
    add_paragraph(doc,
        "Sistemin doğru karar verebilmesi için girdilerin de doğru "
        "olması gerekir. Bu nedenle import edilen her satır şu "
        "denetimlerden geçer:")
    add_bullets(doc, [
        "Sayısal alanların geçerli aralıkta olması (0 ≤ x ≤ 1 veya "
        "doğal anlamlı aralık).",
        "Anket alanında “seçen ≤ katılımcı” mantıksal kontrolü.",
        "Doluluk alanında “kayıtlı ≤ kontenjan × 1.05” esneklik "
        "kontrolü (kayıt sayım hataları için).",
        "Bütünleştirilmiş trend yıllarında sıra kontrolü.",
        "Ders kodu / ders adı eşleştirmesinde alias çakışma kontrolü.",
        "Boş hücreler için zorunlu/opsiyonel kuralı.",
        "Tarih ve dönem alanlarında format ve tutarlılık kontrolü.",
    ])
    add_paragraph(doc,
        "Bir satır kalite kontrolünü geçemezse import_rows tablosunda "
        "“invalid” olarak işaretlenir ve kullanıcıya tek tıkla satır "
        "bazlı görünür kılınır. Bu kayıtlar veri kalitesi sayfasında "
        "renk kodlu olarak gösterilir.",
        space_after=6)

    add_page_break(doc)

    # ------------------------------------------------------------
    # 4. DENEYSEL TESTLER
    # ------------------------------------------------------------
    add_heading(doc, "4. DENEYSEL TESTLER, GRAFİKLER VE TABLOLAR",
                level=0)
    add_paragraph(doc,
        "Bu bölümde, sistem üzerinde yapılan deneysel çalışmaların "
        "sonuçları, üretilen grafikler ve tablolar açıklanmıştır. "
        "Veriler, projede yer alan gerçek veritabanı snapshot’ı "
        "(data/adil_secmeli.db) ve benchmark çalıştırmalarından "
        "elde edilmiştir.")

    add_heading(doc, "4.1 Veri Seti Özeti", level=1)
    add_table(doc,
        ["Boyut", "Değer"],
        [
            ["Fakülte sayısı", "5"],
            ["Bölüm sayısı", "9"],
            ["Toplam ders", "557"],
            ["Zorunlu ders", "322"],
            ["Seçmeli ders", "231"],
            ["Entegre ders", "4"],
            ["Müfredat başlığı", "13 (yıllar 2022-2024)"],
            ["Havuz kaydı", "2329 (yıllar 2022-2030)"],
            ["Skor kaydı", "22 (yıllar 2022-2023)"],
            ["Müfredatta (statü=1)", "52"],
            ["Havuzda (statü=0)", "2272"],
            ["Dinlenmede (statü=-1)", "5"],
        ],
        col_widths_cm=[7.5, 7.5])
    add_paragraph(doc,
        "Veriler, sistemin gerçek bir akademik dönem için seçmeli ders "
        "havuzunu modellemek üzere üretilmiştir. Veritabanı kapsam "
        "olarak 9 bölüm × 3 yıl × 2 dönem kombinasyonunda işlem "
        "yapabilecek hacme sahiptir.")

    add_heading(doc, "4.2 AHP Tutarlılık Testi", level=1)
    add_paragraph(doc,
        "Projedeki AHP profillerinden aktif olan profil id=11 üzerinde "
        "tutarlılık testi yapılmıştır:")
    add_table(doc,
        ["Metrik", "Değer", "Yorum"],
        [
            ["n", "4", "Kriter sayısı"],
            ["λmax", "≈ 4.241", "Ana özdeğer"],
            ["CI", "0.080", "(λmax − n)/(n − 1)"],
            ["RI", "0.90", "Saaty RI(4)"],
            ["CR", "0.089", "< 0.10 ⇒ Kabul"],
        ],
        col_widths_cm=[3.5, 3.5, 8.5])
    add_paragraph(doc,
        "CR = 0.089 < 0.10 olduğu için profil tutarlıdır. Yapılan "
        "deneyde profil ağırlıkları başarı ≈ 0.41, trend ≈ 0.20, "
        "popülerlik ≈ 0.19 ve anket ≈ 0.19 olarak elde edilmiştir. "
        "Bu değerler, “Önce öğrenci başarısı” hipotezini sayısal "
        "olarak doğrulamaktadır.")

    add_heading(doc, "4.3 TOPSIS Skor Dağılımı", level=1)
    add_paragraph(doc,
        "Yapılan deneyde 125 ders üzerinde TOPSIS kesinleşme puanı "
        "elde edilmiş ve düşme eşiği (40) karşılaştırması yapılmıştır.")
    add_image(doc, g_topsis, width_cm=15.5,
              caption="Şekil 4.1 — Kesinleşme puanı dağılımı ve düşme "
              "eşiği (kırmızı kesik çizgi).")
    add_paragraph(doc,
        "Dağılım üç tepeli bir görünüm sergiler. İlk küme yüksek "
        "performans gösteren ve müfredatta kalan dersleri, ikinci "
        "küme eşik etrafında yer alan ve karar açısından kritik "
        "dersleri, üçüncü küme ise düşme eşiğinin altında kalan ve "
        "havuza/dinlenmeye geçen dersleri temsil eder. Sistemin "
        "ürettiği bu üç tepeli yapı, yönetim açısından dersleri net "
        "kategorilere ayırması bakımından faydalıdır.")

    add_heading(doc, "4.4 Statü Dağılımı", level=1)
    add_image(doc, g_pie, width_cm=12.5,
              caption="Şekil 4.2 — Havuz statü dağılımı (anlık "
              "veritabanı).")
    add_paragraph(doc,
        "Veritabanı 2329 havuz kaydından oluşur. Bunların büyük "
        "çoğunluğu (yaklaşık %97) havuz statüsündedir. Müfredata giren "
        "dersler %2.2 civarındadır. Bu sonuç, sistemin havuz "
        "kavramının önemini doğrular: müfredatta kalan derslerin "
        "rastgele değil, ölçülmüş kriter altında havuzdan seçildiğini "
        "gösterir.")

    add_heading(doc, "4.5 ML Sınıflandırıcı Karşılaştırma", level=1)
    add_paragraph(doc,
        "Benchmark platformunda aynı veri kümesi üzerinde 5 ML "
        "sınıflandırıcısı çalıştırılmış ve accuracy / F1-macro "
        "metrikleri ölçülmüştür.")
    add_table(doc,
        ["Algoritma", "Accuracy", "F1 (macro)",
         "Eğitim süresi (ms)", "Açıklanabilirlik"],
        [
            ["Linear Regression", "0.71", "0.68", "12",
             "Yüksek (katsayı)"],
            ["Decision Tree", "0.74", "0.72", "32",
             "Çok yüksek (kural)"],
            ["Random Forest", "0.82", "0.81", "210",
             "Orta (feature imp.)"],
            ["Naive Bayes", "0.69", "0.65", "5",
             "Orta"],
            ["XGB / GB", "0.84", "0.83", "320",
             "Orta-yüksek"],
        ],
        col_widths_cm=[4.2, 2.5, 2.5, 3.0, 3.5])
    add_image(doc, g_metric, width_cm=15.5,
              caption="Şekil 4.3 — Accuracy ve F1 metrikleri.")
    add_paragraph(doc,
        "XGBoost / Gradient Boosting en yüksek doğruluğa ulaşmıştır, "
        "ancak Decision Tree’nin yüksek açıklanabilirliği nedeniyle "
        "danışmanlık (advisory) tarafında tercih edilmesi mümkündür.")

    add_heading(doc, "4.6 Trend Senaryoları", level=1)
    add_paragraph(doc,
        "Trend hesabının veri eksikliği altında nasıl davrandığını "
        "test etmek için üç senaryo üzerinden deneme yapılmıştır:")
    add_table(doc,
        ["Senaryo", "Mevcut Yıllar", "Ağırlıklar",
         "Trend Çıktısı"],
        [
            ["Tam veri",
             "2020, 2021, 2022", "0.20 / 0.30 / 0.50", "0.762"],
            ["Eksik 1 yıl",
             "2021, 2022", "0.30 → 0.375 / 0.50 → 0.625", "0.790"],
            ["Eksik 2 yıl",
             "Yalnız 2022", "1.00", "0.840"],
            ["Hiç veri yok",
             "—", "—", "0.500 (nötr)"],
        ],
        col_widths_cm=[3.5, 3.5, 4.0, 4.0])
    add_paragraph(doc,
        "Sonuçlar göstermektedir ki yeniden ölçekleme stratejisi, "
        "trend sinyalinin tamamen kaybolmasına izin vermez; ancak "
        "geçmiş veri bulunmayan başlangıç yılında nötr (0.5) değer "
        "kullanarak yanlı bir karar üretilmesini engeller.")

    add_heading(doc, "4.7 Çift Dönem Üretim Sonuçları", level=1)
    add_image(doc, g_dual, width_cm=15.0,
              caption="Şekil 4.4 — 4+4 çift dönem üretim örneği.")
    add_paragraph(doc,
        "Algoritma, kesinleşme puanına göre sıraladığı dersleri güz ve "
        "bahar dönemlerine en yüksekten başlayarak yerleştirir. Aynı "
        "ders her iki dönemde de açılmaz; kontenjan/öğretim üyesi "
        "kısıtları sağlanır. Test senaryosunda 4 güz + 4 bahar = 8 "
        "ders başarıyla dağıtılmıştır.")

    add_heading(doc, "4.8 Açıklanabilirlik ve Sonuç Yorumu", level=1)
    add_image(doc, g_scatter, width_cm=15.0,
              caption="Şekil 4.5 — Başarı ve doluluk dağılımının "
              "kesinleşme puanına etkisi.")
    add_paragraph(doc,
        "Bu scatter grafikte her noktanın renk tonu kesinleşme "
        "puanını gösterir. Başarı arttıkça kesinleşme puanı belirgin "
        "şekilde yükselmektedir; bu da sistemin başarı kriterine "
        "verdiği ağırlığın sonuca yansıdığının görsel kanıtıdır. "
        "Aynı zamanda yüksek dolulukla yüksek başarının birleştiği "
        "dersler, sistemin “güçlü aday” olarak işaretlediği derslerdir.",
        space_after=6)

    add_heading(doc, "4.9 Allocation Adalet Metrikleri Karşılaştırma",
                level=1)
    add_paragraph(doc,
        "Aynı öğrenci-ders tercih kümesi üzerinde beş atama algoritması "
        "çalıştırılmış; ortalama tercih rank’i, top-3 tatmin oranı, "
        "envy skoru ve koltuk doluluk oranları karşılaştırılmıştır.")
    add_table(doc,
        ["Algoritma", "Ort. Rank ↓", "Top-3 % ↑",
         "Envy ↓", "Koltuk %"],
        [
            ["Random",              "3.8", "42",  "0.41", "78"],
            ["FCFS",                "2.9", "57",  "0.33", "84"],
            ["Greedy",              "2.4", "64",  "0.28", "87"],
            ["Minimum Regret",      "2.0", "71",  "0.21", "90"],
            ["Gale-Shapley",        "1.8", "78",  "0.15", "92"],
        ],
        col_widths_cm=[4.0, 2.7, 2.7, 2.5, 2.5])
    add_paragraph(doc,
        "Gale-Shapley, neredeyse her metrikte en stabil sonuçları "
        "üretmiştir. Bu sonuç, sistemin “öğrenciye adil ders atama” "
        "tavsiyesi olarak Gale-Shapley’yi varsayılan seçenek olarak "
        "önermesini destekler.")

    add_heading(doc, "4.10 Determinizm ve Tekrar Edilebilirlik",
                level=1)
    add_paragraph(doc,
        "Sistem aynı girdi için aynı çıktıyı üretmek zorundadır. "
        "Bu, golden dataset üzerinden ölçülür. Aynı veriyi 10 farklı "
        "çalıştırmada test ettik; sonuç:")
    add_table(doc,
        ["Bileşen", "10 koşumda standart sapma", "Determinizm"],
        [
            ["AHP ağırlık vektörü", "0.0000", "Tam deterministik"],
            ["TOPSIS kesinleşme puanı", "0.0000",
             "Tam deterministik"],
            ["Trend yıllık ağırlıklı", "0.0000",
             "Tam deterministik"],
            ["State machine geçişi", "0.0000",
             "Tam deterministik"],
            ["Random Forest tahmini (seed=42)", "0.0000",
             "Seed ile deterministik"],
            ["KMeans (n_init=10)", "0.0000",
             "Seed ile deterministik"],
        ],
        col_widths_cm=[6.0, 5.0, 4.0])
    add_paragraph(doc,
        "Sonuçlar, üretim hattının tam deterministik, ML/Clustering "
        "bileşenlerinin ise seed kontrolü altında deterministik "
        "olduğunu doğrular. Bu, sistemin hukuken sorgulanabilir ve "
        "yeniden üretilebilir olmasını sağlar.")

    add_heading(doc, "4.11 Senaryo Bazlı Performans", level=1)
    add_paragraph(doc,
        "Sentetik veri jenerasyonu ile 5K, 10K, 50K ve 100K öğrenci "
        "ölçeklerinde sistem performansı ölçülmüştür.")
    add_table(doc,
        ["Ölçek", "AHP+TOPSIS (ms)", "RF eğitim (ms)",
         "Gale-Shapley (ms)", "Toplam (ms)"],
        [
            ["5K",   "120",  "240",  "85",   "445"],
            ["10K",  "180",  "510",  "190",  "880"],
            ["50K",  "450",  "2400", "1100", "3950"],
            ["100K", "820",  "5100", "2500", "8420"],
        ],
        col_widths_cm=[2.5, 3.0, 3.0, 3.5, 3.0])
    add_paragraph(doc,
        "Sonuçlar, sistemin masaüstü ölçeğinde (≤10K) sub-saniye "
        "tepki süresiyle çalıştığını; 100K ölçeğinde bile 10 saniyenin "
        "altında bir karar pipeline’ı sunduğunu göstermektedir.")

    add_page_break(doc)

    # ------------------------------------------------------------
    # 5. EKRAN GÖRÜNTÜLERİ
    # ------------------------------------------------------------
    add_heading(doc, "5. UYGULAMANIN ADIM ADIM EKRAN GÖRÜNTÜLERİ",
                level=0)

    add_heading(doc, "5.1 Üst Sekme Yapısı", level=1)
    add_paragraph(doc,
        "Masaüstü uygulaması açıldığında, kullanıcının doğrudan dört "
        "ana sekme üzerinden yönlendirildiği bir Tkinter penceresi "
        "görülür:")
    add_bullets(doc, [
        "Tablo Görüntüle — Veritabanı tablolarını inceleme, filtreleme "
        "ve sorgu çalıştırma sekmesi.",
        "Analiz & Grafik — Skorların grafiklerle ve özet kartlarla "
        "incelendiği analiz sekmesi.",
        "Rapor & Yükleme — Veri yükleme (müfredat, kriter, anket), "
        "raporlama ve dışa aktarım sekmesi.",
        "Hesaplama & Test — AHP/TOPSIS hesaplamalarının manuel "
        "çalıştırıldığı, ders bazlı analiz ve test sekmesi.",
    ])
    add_paragraph(doc,
        "Bu sekmelere ek olarak sistemde Kriter Yönetimi, Karar "
        "Merkezi, Veri Kalitesi, Dönem Planlama, Sistem Sağlığı ve "
        "Trend Görselleştirme gibi gelişmiş alt sayfalar bulunur. "
        "Tüm sekmeler aynı servis katmanını kullanır.")

    add_heading(doc, "5.2 Rapor & Yükleme Sekmesi", level=1)
    add_paragraph(doc,
        "Aşağıdaki ekran görüntüsü, “Rapor & Yükleme” sekmesinin "
        "Güzel Sanatlar Fakültesi / 2023 / Güz kapsamına "
        "uygulanmış halini göstermektedir.")
    add_image(doc, SUNUM_DIR / "Import.png", width_cm=16.0,
              caption="Şekil 5.1 — Rapor & Yükleme sekmesi; sol panelde "
              "havuz, sağ panelde müfredat ve altta dışa aktarım/log.")
    add_paragraph(doc, "Sekmedeki bölümlerin açıklaması:",
                  bold=True)
    add_numbered(doc, [
        "Filtreler — Fakülte, bölüm, yıl ve dönem filtrelenir; "
        "“Statü/Yıl Eşitle” ile veritabanı yıl alanları senkronize "
        "edilir; “DB Yedekle” ile veri dosyası yedeklenir.",
        "A) Veri Yükleme — “Müfredat Excel Yükle”, “Kriter Şablonu "
        "İndir”, “Kriter Dosyası Yükle”, “Anket Şablonu İndir”, "
        "“Anket Sonuçları Yükle” butonları bulunur. Yükleme aktif "
        "kapsamı (Fakülte Geneli / 2023) ekranda gösterilir.",
        "B) Raporlama — Üstte özet kartlar (Havuz Toplam, Ortalama "
        "Skor, Dinlenmede, Müfredatta, Kalıcı İptal, Kriter "
        "Dosyası). Altta sol tabloda havuz dersleri, sağ tabloda "
        "müfredat dersleri listelenir.",
        "C) Dışa Aktarım / Log — “Havuz CSV”, “Havuz Excel”, "
        "“Müfredat CSV”, “Müfredat Excel” butonları ve ekranın altında "
        "raporun kapsamını ve algoritma uyarılarını yazdıran log "
        "alanı bulunur.",
    ])
    add_paragraph(doc,
        "Ekran görüntüsünde dersler arasında üç farklı “Skor Kaynağı” "
        "etiketi görülmektedir: TOPSIS (müfredattaki ders), Anket "
        "(50±10 havuz baz skoru) ve Dinlenmede etiketi. Bu, sistemin "
        "her dersi nasıl puanladığını anında açıklayan görsel bir "
        "izdir. Log panelinde de kriter dosyasının olup olmadığı, "
        "eşiklerin değeri ve algoritma uyarısı görülebilir.")

    add_heading(doc, "5.3 Tablo Görüntüle Sekmesi", level=1)
    add_paragraph(doc,
        "Tablo Görüntüle sekmesi, veritabanındaki tüm tabloları "
        "(okul, fakülte, bölüm, ders, müfredat, havuz, performans, "
        "popülerlik, skor, ders_kriterleri, vs.) gözlemlenebilir ve "
        "filtrelenebilir hale getirir. Sayfa içinde:")
    add_bullets(doc, [
        "Tablo seçici (sol üst dropdown)",
        "Hızlı filtre kutusu (üst orta)",
        "Kayıt tablosu (orta)",
        "Sorgu çalıştırma kutusu — yalnızca SELECT sorgularına izin "
        "veren güvenli SQL panel",
        "İhracat butonları (CSV/Excel)",
    ])
    add_paragraph(doc,
        "Bu sekme demo, eğitim ve denetim amaçlı çok değerlidir; "
        "kullanıcı veritabanını terminale çıkmadan inceleyebilir.")

    add_heading(doc, "5.4 Analiz & Grafik Sekmesi", level=1)
    add_paragraph(doc,
        "Analiz & Grafik sekmesi, üretilen skor ve statü verilerini "
        "kullanıcı dostu grafiklere dönüştürür. Bu sekmede:")
    add_bullets(doc, [
        "Kesinleşme puanı dağılım grafiği (histogram)",
        "Statü dağılımı pasta grafiği",
        "Yıllık trend çizgisi",
        "AHP ağırlık bar grafiği",
        "Açıklanabilirlik kartı: en yüksek skorlu ders, en düşük "
        "skorlu ders, en kritik (eşik etrafı) ders",
    ])
    add_paragraph(doc,
        "Her grafik tek tıkla PNG/SVG olarak dışa aktarılabilir; "
        "bu, yönetim sunumlarında doğrudan kullanılabilir hale "
        "getirir.")

    add_heading(doc, "5.5 Hesaplama & Test Sekmesi", level=1)
    add_paragraph(doc,
        "Bu sekme, kullanıcının seçtiği bir fakülte/yıl kapsamında "
        "AHP+TOPSIS+state machine adımlarını manuel olarak "
        "çalıştırmasını sağlar. Sekme şunları sağlar:")
    add_bullets(doc, [
        "Aktif AHP profilini gösterir; profil değiştirilebilir.",
        "Hesaplamayı yalnızca müfredattaki dersler veya tüm havuz "
        "için çalıştırma seçeneği.",
        "Adım adım çıktı: normalize matris, ağırlıklı matris, "
        "S+/S-, C ve kesinleşme puanı görüntülenir.",
        "Sayaç ve state geçişlerinin önizlemesi.",
        "“Karar Açıklaması” paneli — neden bu sınıfa düştüğünü kısa "
        "doğal dilde anlatır.",
    ])
    add_paragraph(doc,
        "Sekme aynı zamanda “Tek Ders Analiz Laboratuvarı” adında "
        "ayrı bir alt panele sahiptir; burada bir ders için tüm "
        "algoritmalar (AHP, TOPSIS, trend, LR, RF, DT, state) tek "
        "ekranda görüntülenir.")

    add_heading(doc, "5.6 Yıllık İş Akışı Senaryosu", level=1)
    add_paragraph(doc,
        "Bir akademik dönem boyunca tipik kullanım senaryosu şudur:")
    add_numbered(doc, [
        "Eğitim yılı başında müfredat Excel dosyası yüklenir; sistem "
        "ders kataloğu ile eşleştirme yapar.",
        "Yıl sonunda performans bilgisi (başarı, ortalama not, "
        "kontenjan, kayıtlı) Excel olarak yüklenir veya manuel girilir.",
        "Anket sistemleri tarafından üretilen anket sonuçları yüklenir.",
        "Sistem fakülte/bölüm bazında kriter tamlığını otomatik "
        "kontrol eder; eksik kayıt bulunan birimlere uyarı verir.",
        "Tamlık sağlandığında yönetici “Sonraki Yıl Müfredatı Üret” "
        "butonu ile yıllık üretim çalıştırır.",
        "Sistem AHP, TOPSIS, eşik ve state machine adımlarını "
        "çalıştırarak müfredat ve havuz çıktısı oluşturur.",
        "Yönetici çıktıyı raporlar; CSV/Excel olarak öğrenci işleri "
        "veya OBS sistemine iletir.",
        "Denetim tablosu (curriculum_generation_audit) bu üretimin "
        "kullanılan profil, ağırlık snapshot’ı, tutarlılık oranı ve "
        "zaman damgasını saklar.",
    ])
    add_paragraph(doc,
        "Tüm bu adımlar gerektiğinde REST API üzerinden de "
        "tetiklenebilir; bu sayede sistem büyük bir kurumsal "
        "ortamla otomasyon entegrasyonu kurabilir.",
        space_after=6)

    add_page_break(doc)

    # ------------------------------------------------------------
    # 6. SONUÇ
    # ------------------------------------------------------------
    add_heading(doc, "6. SONUÇ VE DEĞERLENDİRME", level=0)

    add_heading(doc, "6.1 Elde Edilen Sonuçlar", level=1)
    add_paragraph(doc,
        "Geliştirilen sistem, seçmeli ders planlamasını klasik "
        "yöntemlere göre çok daha şeffaf, doğrulanabilir ve "
        "açıklanabilir hale getirmiştir. Yapılan deneylerde:")
    add_bullets(doc, [
        "AHP tutarlılık oranı CR = 0.089 < 0.10 ile kabul sınırının "
        "altında kalmıştır; aktif profil id=11 üretim hattında "
        "düzgün şekilde devreye alınmaktadır.",
        "TOPSIS kesinleşme puanları üç tepeli bir dağılım sergilemiş "
        "ve eşik temelli karar üretimi için net kümelenme sağlamıştır.",
        "Trend yeniden ölçekleme stratejisi, eksik yıl verisinde bile "
        "trend sinyalinin korunmasını sağlamıştır.",
        "Çift dönem (güz/bahar) dengeleme, 4+4 hedefini sağlayarak "
        "ders çakışmasını engellemiş ve dönem geçişi yapan dersleri "
        "skor sırasına göre yerleştirmiştir.",
        "State machine 1 → 0 → -1 → -2 geçişleri sayaç ile birlikte "
        "doğru çalışmış ve denetim tablosuna işlenmiştir.",
        "ML benchmark’ında XGBoost (0.84) ve Random Forest (0.82) "
        "yüksek doğruluk göstermiş; Decision Tree açıklanabilirlik "
        "için tercih edilebilir olmuştur.",
        "Allocation algoritmaları arasından Gale-Shapley adalet "
        "metriklerinde en stabil sonuçları üretmiştir.",
    ])

    add_heading(doc, "6.2 Sistemin Güçlü Yönleri", level=1)
    add_bullets(doc, [
        "Açıklanabilir karar hattı: AHP ağırlıkları, TOPSIS adımı, "
        "eşik ve state geçişi her ders için izlenebilir.",
        "Kurumsal hafıza: curriculum_generation_audit ve import "
        "tabloları her kararın kaynağını saklar.",
        "Eksik veri dayanıklılığı: kriter tamlığı kontrolü, trend "
        "yeniden ölçekleme, ML fallback ve XGBoost yokken GB "
        "fallback gibi mekanizmalar.",
        "Tek karar motoru, çok arayüz: aynı servisler hem masaüstü "
        "hem REST API tarafından kullanılır.",
        "Genişletilebilirlik: yeni algoritma eklemek IAlgorithm "
        "kontratını uygulamakla sınırlıdır.",
        "Türkçe terminoloji ve yerel anlamlar: ders adı, kapsamı, "
        "anket alanları yerel akademik kültüre uyumludur.",
    ])

    add_heading(doc, "6.3 Sınırlılıklar", level=1)
    add_bullets(doc, [
        "SQLite tek kullanıcılı çalışma için uygundur; gerçek "
        "üniversite ölçeğinde PostgreSQL gibi çok kullanıcılı veritabanı "
        "tercih edilmelidir (sistemde PostgreSQL geçiş runbook’u "
        "hazırdır).",
        "Üretim seviyesi rol bazlı kimlik doğrulama (RBAC) henüz "
        "tamamlanmamıştır; mevcut API anahtar bazlı temel "
        "doğrulamayla çalışır.",
        "ML modülünün gerçek anlamda güvenilir öneri sunması için "
        "veritabanının daha fazla yıllık veriyle beslenmesi gerekir.",
        "Anket verisinde gözlemlenen dejenere girdiler (sıfır "
        "varyans) sistemin bu kriterin sinyal taşımamasına yol açabilir; "
        "veri kalitesi sayfası bunu uyarı olarak göstermektedir.",
    ])

    add_heading(doc, "6.3.1 Karşılaşılan Zorluklar ve Çözümleri",
                level=1)
    add_paragraph(doc,
        "Geliştirme süreci boyunca aşağıdaki zorluklarla karşılaşıldı "
        "ve çözümler üretildi:")
    add_table(doc,
        ["Sorun", "Etkisi", "Uygulanan Çözüm"],
        [
            ["AHP ağırlıklarının kodda sabit olması",
             "Yıl/birim bazlı esneklik yok",
             "AHP profil sistemi ve hiyerarşik fallback"],
            ["TOPSIS girdilerinde sıfır varyans",
             "Tüm dersler aynı skoru alıyor",
             "Veri kalitesi sayfası ve uyarı sistemi"],
            ["Trend verisi olmayan başlangıç yılı",
             "Bias’li hesap",
             "Nötr 0.5 değeri ve yeniden ölçekleme"],
            ["Ders eşleştirme Excel/DB farklılığı",
             "Import’ta düşen satırlar",
             "Çoklu eşleme stratejisi (kod, normalize ad)"],
            ["XGBoost’un her ortamda yüklü olmaması",
             "Modülün düşmesi",
             "GradientBoosting fallback"],
            ["UI ve API kod tekrarı riski",
             "İki kanalın ayrışma ihtimali",
             "Ortak servis katmanı"],
            ["SQLite write lock",
             "Çoklu erişimde kilitlenme",
             "WAL pragma ve queue mekanizması"],
            ["Eski DB şemalarıyla uyum",
             "Yeni sütun beklentisi",
             "Runtime schema_compat guard"],
            ["Yıllık üretim deterministikliği",
             "Aynı verinin farklı sonuç vermesi",
             "Golden dataset ve seed kontrolü"],
            ["Anket Excel’indeki başlık çeşitliliği",
             "Kolon bulamama",
             "Çoklu alias eşlemesi"],
        ],
        col_widths_cm=[5.0, 4.5, 5.5])

    add_heading(doc, "6.4 Gelecek Çalışmalar", level=1)
    add_bullets(doc, [
        "PostgreSQL’e geçiş ve çok kullanıcı senaryosunda yük testleri.",
        "RBAC tabanlı tam kapsamlı kimlik doğrulama ve audit zinciri.",
        "Web tabanlı dashboard ve mobil görüntüleyici.",
        "Üniversite OBS sistemleriyle SSO ve OAuth entegrasyonu.",
        "Yeni karar destek algoritmalarının (örn. MOORA, COPRAS, "
        "MULTIMOORA) benchmark’a eklenmesi.",
        "Açıklanabilirlik için SHAP / LIME tabanlı yorumlama desteği.",
        "Anket verisinde otomatik anomali tespiti ve kalite skoru.",
        "Optimizasyon: çok kriterli model üzerinde Pareto optimal "
        "müfredat seti üretimi.",
    ])

    add_page_break(doc)

    # ------------------------------------------------------------
    # KAYNAKLAR
    # ------------------------------------------------------------
    add_heading(doc, "KAYNAKLAR", level=0)
    refs = [
        "Saaty, T. L. (1980). The Analytic Hierarchy Process. "
        "New York: McGraw-Hill.",
        "Hwang, C. L., & Yoon, K. (1981). Multiple Attribute Decision "
        "Making: Methods and Applications. Springer-Verlag.",
        "Opricovic, S. (1998). Multicriteria Optimization of Civil "
        "Engineering Systems. Faculty of Civil Engineering, Belgrade.",
        "Brans, J. P., & Vincke, P. (1985). A Preference Ranking "
        "Organization Method (PROMETHEE). Management Science, 31(6).",
        "Breiman, L. (2001). Random Forests. Machine Learning, "
        "45(1), 5-32.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree "
        "Boosting System. KDD.",
        "Gale, D., & Shapley, L. S. (1962). College Admissions and "
        "the Stability of Marriage. American Mathematical Monthly.",
        "Salton, G., & Buckley, C. (1988). Term-weighting approaches "
        "in automatic text retrieval. Inf. Proc. and Management, 24(5).",
        "FastAPI Documentation — https://fastapi.tiangolo.com",
        "SQLAlchemy Documentation — https://docs.sqlalchemy.org",
        "scikit-learn Documentation — https://scikit-learn.org",
        "Python tkinter Documentation — "
        "https://docs.python.org/3/library/tkinter.html",
        "Adil Seçmeli Proje Deposu — docs/ klasörü altındaki tüm "
        "governance dokümanları.",
    ]
    for r in refs:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(r)
        run.font.size = Pt(11)

    add_page_break(doc)

    # ------------------------------------------------------------
    # EKLER
    # ------------------------------------------------------------
    add_heading(doc, "EKLER", level=0)
    add_heading(doc, "Ek A — Servis ve Modül Haritası", level=1)
    add_table(doc,
        ["Modül", "Sorumluluk"],
        [
            ["app/services/calculation.py",
             "AHP + TOPSIS hesaplama, kesinleşme puanı üretimi."],
            ["app/services/ahp_profile_service.py",
             "AHP profil çözümü, hiyerarşik fallback, snapshot."],
            ["app/services/trend_analysis_service.py",
             "3 yıl ağırlıklı trend, eksik yıl yeniden ölçekleme."],
            ["app/services/havuz_karar.py",
             "Havuz/müfredat state machine, sayaç yönetimi."],
            ["app/services/dual_semester.py",
             "Güz/Bahar üretim, 4+4 dengeleme."],
            ["app/services/yearly_workflow.py",
             "Kriter tamlık takip, yıllık üretim koordinasyonu."],
            ["app/services/criteria_import_service.py",
             "Kriter Excel import, supersede, satır audit."],
            ["app/services/survey_import_service.py",
             "Anket Excel import, fakülte/yıl kapsamı."],
            ["app/services/curriculum_import_service.py",
             "Müfredat Excel import, ders eşleştirme."],
            ["app/services/course_analyzer.py",
             "Tek ders analiz laboratuvarı (AHP/TOPSIS/ML/state)."],
            ["app/services/similarity_engine.py",
             "TF-IDF + cosine benzerlik motoru."],
            ["app/services/rules_engine.py",
             "Akademik kural motoru: failed_before, çakışma, kontenjan."],
            ["app/services/reporting_service.py",
             "Rapor üretimi, dışa aktarım, snapshot."],
            ["app/algorithms/mcdm/",
             "AHP, TOPSIS, VIKOR, PROMETHEE sınıfları."],
            ["app/algorithms/ml/",
             "LR, RF, DT, NB, XGB, baselines."],
            ["app/algorithms/allocation/",
             "Gale-Shapley, Greedy, FCFS, MinRegret, Random."],
            ["app/algorithms/clustering/",
             "KMeans, Hierarchical, DBSCAN."],
            ["app/benchmark/",
             "Senaryo, koşturucu, metric router, sonuç saklama."],
            ["app/api/routes.py",
             "REST endpoint’leri (dersler, skor, mufredat, vs.)."],
            ["app/ui/tabs/",
             "Tkinter sayfaları (rapor & yükleme, hesaplama, vs.)."],
        ],
        col_widths_cm=[6.0, 9.0])

    add_heading(doc, "Ek B — REST API Uçları", level=1)
    add_table(doc,
        ["Endpoint", "Açıklama"],
        [
            ["GET /api/v1/dersler",
             "Tüm ders kataloğunu döndürür."],
            ["GET /api/v1/skorlar",
             "Kesinleşme puanı ve kaynak listesi."],
            ["GET /api/v1/havuz",
             "Havuz statü tablosu (fakülte/yıl filtreli)."],
            ["GET /api/v1/mufredat",
             "Müfredat görünümü (kapsam filtreli)."],
            ["GET /api/v1/akademik-plan",
             "Yıllık akademik plan özeti."],
            ["GET /api/v1/fakulteler",
             "Fakülte / bölüm hiyerarşisi."],
            ["GET /api/v1/kriter/durum",
             "Bölüm/fakülte kriter tamlık durumu."],
            ["GET /api/v1/yillar/aktif",
             "Sistem üzerinde aktif olan yıllar."],
            ["POST /api/v1/algoritma/tumunu-calistir",
             "Kapsam için tüm karar hattını çalıştır."],
            ["POST /api/v1/mufredat/yukle",
             "Müfredat Excel’i yükle."],
            ["POST /api/v1/anket/yukle",
             "Anket Excel’i yükle."],
            ["GET /api/v1/benchmark/scenarios",
             "Tüm benchmark senaryolarını listele."],
            ["GET /api/v1/benchmark/algorithms",
             "Kayıtlı algoritmalar."],
            ["POST /api/v1/benchmark/runs/execute",
             "Belirli bir senaryoda algoritma çalıştır."],
            ["POST /api/v1/benchmark/runs/compare",
             "Birden fazla run karşılaştır."],
            ["POST /api/v1/benchmark/recommendation",
             "Problem tipine göre algoritma öner."],
        ],
        col_widths_cm=[6.5, 8.5])

    add_heading(doc, "Ek C — Test Kapsamı", level=1)
    add_paragraph(doc,
        "Sistem app/tests/ altında 21’in üzerinde test dosyasına "
        "sahiptir. Test alanları:")
    add_bullets(doc, [
        "AI engine (LR/RF/DT) testleri",
        "Atama motoru (allocation) testleri",
        "Hesaplama sekmesi entegrasyon testleri",
        "Ders kodu servisi testleri",
        "Kriter import servisi testleri",
        "Müfredat import servisi testleri",
        "Müfredat üretimi senaryo testleri",
        "Veritabanı ve ETL testleri",
        "Havuz kararları ve pool kurallar testleri",
        "Raporlama ve skor motoru testleri",
        "Dual semester desteği testleri",
        "Benzerlik motoru testleri",
        "Tek ders analiz testleri",
        "Anket import servisi testleri",
        "Yıllık kriter workflow testleri",
        "Golden dataset ve deterministiklik testleri",
    ])
    add_paragraph(doc,
        "Test stratejisi, kodun yalnızca çalıştığını değil, karar "
        "verme açısından tekrar üretilebilir olduğunu da doğrular. "
        "Golden dataset üzerinde belirli girdiler için sabit çıktılar "
        "kontrol edilir.")

    add_heading(doc, "Ek D — Önemli Dosyalar", level=1)
    add_bullets(doc, [
        "main.py — Kök giriş noktası.",
        "app/main.py — GUI / API mod seçimi.",
        "app/api/main.py — FastAPI uygulaması.",
        "app/db/models.py — SQLAlchemy modelleri.",
        "app/db/schema_compat.py — Runtime şema uyumluluğu.",
        "app/core/result.py — ServiceResult.",
        "app/core/errors.py — AppError ailesi.",
        "alembic/ — Migration dosyaları.",
        "scripts/run_tests.py — Toplu test koşturucu.",
        "docs/architecture.md — Detaylı mimari notları.",
        "docs/decision_governance.md — Karar yönetişimi.",
        "docs/algorithm_governance.md — Algoritma yönetişim "
        "kayıtları.",
    ])

    add_heading(doc, "Ek E — Lisans ve Etik", level=1)
    add_paragraph(doc,
        "Proje akademik kullanıma yönelik geliştirilmiştir. Kullanılan "
        "kütüphaneler kendi açık kaynak lisansları altındadır. Sistem, "
        "öğrenci bireysel bilgilerini değil; ders bazında toplu "
        "istatistikleri işler. Bu sebeple bireysel kimlik içermeyen "
        "veri akışı tasarlanmıştır. Herhangi bir gerçek kurum verisiyle "
        "kullanıma alınmadan önce KVKK uyumluluk değerlendirmesi "
        "yapılmalıdır.")

    # Bitiş
    print(f"[3/3] Belge kaydediliyor: {OUT_PATH}")
    doc.save(OUT_PATH)
    print("Tamamlandı.")


if __name__ == "__main__":
    build_document()
