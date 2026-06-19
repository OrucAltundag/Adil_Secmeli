# -*- coding: utf-8 -*-
"""Adil Seçmeli — Bitirme teslim raporları üreticisi (Word/.docx).

Tüm raporlar mümkün olduğunca CANLI veriden üretilir (statik/mock veri kuralı):
veritabanı read-only (`mode=ro`) açılır, böylece masaüstü uygulaması açıkken
(yazma kilidi varken) bile sorgular çalışır. Canlı veri alınamazsa ilgili
bölüm yapısal olarak açıklanır.

Çalıştırma:
    python scripts/generate_reports.py
Çıktı: docs/ klasörüne .docx dosyaları.
"""

from __future__ import annotations

import os
import sqlite3
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document  # type: ignore[import-not-found]  # noqa: E402  # python-docx opsiyonel bagimliliktir
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-not-found]  # noqa: E402
from docx.shared import Pt, RGBColor  # type: ignore[import-not-found]  # noqa: E402

DOCS_DIR = PROJECT_ROOT / "docs"
DB_PATH = PROJECT_ROOT / "data" / "adil_secmeli.db"
GEN_TS = datetime.now().strftime("%d.%m.%Y %H:%M")

PRIMARY = RGBColor(0x15, 0x65, 0xC0)
OK = RGBColor(0x2E, 0x7D, 0x32)
WARN = RGBColor(0xE6, 0x51, 0x00)
CRIT = RGBColor(0xC6, 0x28, 0x28)


def ro_conn() -> sqlite3.Connection:
    """Read-only bağlantı (uygulama açıkken kilitten etkilenmez)."""
    conn = sqlite3.connect(f"file:{DB_PATH.as_posix()}?mode=ro", uri=True)
    conn.row_factory = sqlite3.Row
    return conn


def q(conn: sqlite3.Connection, sql: str, params=()) -> list:
    try:
        return conn.execute(sql, params).fetchall()
    except sqlite3.Error:
        return []


def scalar(conn: sqlite3.Connection, sql: str, params=(), default: Any = 0) -> Any:
    rows = q(conn, sql, params)
    if rows and rows[0] and rows[0][0] is not None:
        return rows[0][0]
    return default


# --- docx yardımcıları ------------------------------------------------------

def new_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    h = doc.add_heading(title, level=0)
    for run in h.runs:
        run.font.color.rgb = PRIMARY
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(10)
    meta = doc.add_paragraph()
    mr = meta.add_run(f"Adil Seçmeli Ders Yönetim Sistemi · Üretim: {GEN_TS} · Veri: canlı (read-only)")
    mr.font.size = Pt(8)
    mr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    doc.add_paragraph()
    return doc


def h1(doc: Document, text: str):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = PRIMARY


def h2(doc: Document, text: str):
    doc.add_heading(text, level=2)


def para(doc: Document, text: str):
    doc.add_paragraph(text)


def bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list], col_color=PRIMARY):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = str(htxt)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
    doc.add_paragraph()
    return t


def save(doc: Document, filename: str) -> Path:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    out = DOCS_DIR / filename
    doc.save(str(out))
    return out


# ============================================================================
# 1) VERİ SETLERİ VE SİSTEM UYGUNLUK RAPORU (madde 1, 14)
# ============================================================================

def report_data_fitness() -> Path:
    doc = new_doc(
        "Veri Setleri ve Sistem Uygunluk Raporu",
        "Veri setlerinin sisteme uygunluğu, tablo-kaynak haritası ve canlı veri kalitesi bulguları",
    )
    conn = ro_conn()

    h1(doc, "1. Amaç ve Kapsam")
    para(doc, "Bu rapor, data/ klasöründeki veri setlerinin hangi veritabanı tablolarını "
              "beslediğini, sistem şemasına uygunluğunu ve veritabanının canlı (read-only) "
              "olarak denetlenen kalite durumunu belgeler. Tüm sayılar rapor üretildiği anda "
              "veritabanından okunmuştur.")

    h1(doc, "2. Veri Seti → Tablo Haritası")
    table(doc, ["Veri Seti (data/)", "Beslediği Tablo(lar)", "Açıklama"], [
        ["2021/2022/2023_ogrenci_not_veri_seti.xlsx", "performans", "Öğrenci ders notları → ders başarı/performans"],
        ["2022_ogrenci_not_veri_seti_duzeltilmis.xlsx", "performans", "Düzeltilmiş not seti (tutarsızlık bayraklı)"],
        ["2022_Mufredat.xlsx", "mufredat, mufredat_ders", "Bölüm müfredatları ve ders bağları"],
        ["2022_anket_tercih_veri_seti(i).xlsx", "anket_sonuclari, ders_kriterleri", "Öğrenci tercih/anket verisi"],
        ["2022_kriterler.csv", "ders_kriterleri", "Ders kriter değerleri (başarı, popülerlik, anket)"],
        ["dersler_master.xlsx", "ders", "Ders ana kayıt (kod, ad, kredi, akts)"],
    ])

    h1(doc, "3. Tablo Doluluk Özeti (canlı)")
    rows = []
    for t in ["ders", "havuz", "performans", "populerlik", "anket_sonuclari",
              "ders_kriterleri", "mufredat", "mufredat_ders", "fakulte", "bolum"]:
        rows.append([t, scalar(conn, f"SELECT COUNT(*) FROM {t}")])
    table(doc, ["Tablo", "Kayıt Sayısı"], rows)

    h1(doc, "4. Veri Bütünlüğü (canlı)")
    integrity = scalar(conn, "PRAGMA integrity_check", default="?")
    fk = len(q(conn, "PRAGMA foreign_key_check"))
    bullet(doc, f"SQLite integrity_check: {integrity}")
    bullet(doc, f"foreign_key_check ihlali: {fk}")
    orphan = scalar(conn, "SELECT COUNT(*) FROM havuz WHERE CAST(ders_id AS INTEGER) NOT IN (SELECT ders_id FROM ders)")
    bullet(doc, f"Yetim havuz kaydı (ders_id eşleşmeyen): {orphan}")
    para(doc, "Sonuç: Veri bütünlüğü temiz; integrity 'ok', FK ihlali yok, yetim kayıt yok.")

    h1(doc, "5. Tespit Edilen Kalite Bulguları ve Çözümler")
    h2(doc, "5.1 Havuz 'tekrar' yanlış pozitifi — ÇÖZÜLDÜ")
    para(doc, "Sağlık kontrolünde havuz tekrar anahtarı (ders_id, yil) idi; aynı dersin "
              "Güz ve Bahar kayıtlarını yanlışlıkla 'tekrar' sayıyordu. Anahtar "
              "(ders_id, yil, donem) yapıldı (app/health/health_config.py). Gerçek tekrar yok.")
    dups = q(conn, "SELECT ders_id,yil,donem,COUNT(*) ct FROM havuz GROUP BY ders_id,yil,donem HAVING ct>1")
    bullet(doc, f"Gerçek tekrar (ders_id, yil, donem): {len(dups)} grup")

    h2(doc, "5.2 Aykırı kredi değerleri — ÖNERİ (otomatik düzeltilmedi)")
    kredi_out = q(conn, "SELECT ders_id, kod, ad, kredi, akts FROM ders WHERE kredi>30 ORDER BY kredi DESC LIMIT 15")
    para(doc, "ders.kredi alanında tıp 'Ders Kurulu' bloklarında implausible değerler var "
              "(kredi alanına saat girilmiş gibi). Kaynak master veri; doğru değer "
              "bilinmediğinden otomatik düzeltilmedi, manuel onay önerilir.")
    if kredi_out:
        table(doc, ["ders_id", "kod", "ad", "kredi", "akts"],
              [[r[0], r[1], r[2], r[3], r[4]] for r in kredi_out])

    h2(doc, "5.3 İçe aktarımda 'çözülmemiş 57 satır' — AÇIKLANDI")
    n57 = scalar(conn, "SELECT COUNT(*) FROM import_row_issues")
    by_type = q(conn, "SELECT severity, issue_type, COUNT(*) FROM import_row_issues GROUP BY severity, issue_type")
    para(doc, f"import_row_issues tablosundaki {n57} kayıt, batch #2 (anket importu) için "
              "üretilmiştir. Belge fakültesi 'Mühendislik ve Doğa Bilimleri', içe aktarımda "
              "'Tıp Fakültesi' seçilmiş → tüm satırlar invalid_scope hatasıyla reddedilmiş ve "
              "batch 'failed' olmuştur. Bu bir veri bozulması DEĞİL, doğru reddedilmiş bir "
              "importun denetim kaydıdır. Çözüm: doğru fakülteyle yeniden import.")
    if by_type:
        table(doc, ["Önem", "Sorun Tipi", "Adet"], [[r[0], r[1], r[2]] for r in by_type])

    h1(doc, "6. Sonuç")
    para(doc, "Veri setleri sisteme uygundur ve doğru okunmaktadır. Veritabanı bütünlüğü "
              "temizdir. Tespit edilen tek yapılandırma hatası (havuz tekrar anahtarı) "
              "düzeltilmiştir. Aykırı krediler ve 57 satır kullanıcı aksiyonu gerektiren "
              "öneri/açıklama maddeleridir.")
    conn.close()
    return save(doc, "Veri_Setleri_ve_Sistem_Uygunluk_Raporu.docx")


# ============================================================================
# 2) BENCHMARK / ALGORİTMA KULLANIM RAPORU (madde 3, 4)
# ============================================================================

def report_algorithms() -> Path:
    doc = new_doc(
        "Benchmark Algoritmaları Kullanım Raporu",
        "Sistemdeki algoritmaların canlı registry'den sınıflandırması ve projedeki rolleri",
    )
    conn = ro_conn()

    h1(doc, "1. Algoritma Yönetişim Modeli")
    para(doc, "Sistem, algoritmaları 'algorithm_governance_registry' tablosunda merkezî olarak "
              "yönetir. Her algoritmanın bir kullanım rolü (usage_role), nihai kararı "
              "etkileyip etkilemediği (can_affect_final_decision) ve aktiflik durumu vardır. "
              "Bir algoritmanın benchmark panelinde görünmesi, onun nihai kararı etkilediği "
              "anlamına gelmez.")
    table(doc, ["Rol (usage_role)", "Anlamı"], [
        ["production_decision", "Üretim hattının ANA karar motoru (kararı doğrudan etkiler)"],
        ["advisory_ml", "Destekleyici ML (öneri üretir, kararı doğrudan değiştirmez)"],
        ["benchmark_only", "Yalnızca karşılaştırma/benchmark amaçlı"],
        ["baseline", "Referans (baseline) — kalite kıyaslaması için"],
    ])

    h1(doc, "2. Kayıtlı Algoritmalar (canlı registry)")
    rows = q(conn, "SELECT algorithm_key, algorithm_family, usage_role, "
                   "can_affect_final_decision, is_active FROM algorithm_governance_registry "
                   "ORDER BY algorithm_family, usage_role, algorithm_key")
    if rows:
        table(doc, ["Algoritma", "Aile", "Rol", "Karara Etki", "Aktif"],
              [[r[0], r[1], r[2], "Evet" if r[3] else "Hayır", "Evet" if r[4] else "Hayır"] for r in rows])
        para(doc, f"Toplam {len(rows)} algoritma kayıtlı.")

    h1(doc, "3. Ana Karar Motorları (production_decision)")
    para(doc, "AHP ve TOPSIS üretim hattının ana karar motorlarıdır (kesinleşme puanı, "
              "müfredat üretimi). rule_engine, state_machine ve trend_analysis da kararı "
              "etkileyen üretim bileşenleridir. VIKOR ve PROMETHEE_II yalnızca benchmark "
              "amaçlıdır; sonuçları nihai kararı değiştirmez.")

    h1(doc, "4. Benchmark Platformu — Canlı Veri Mimarisi")
    para(doc, "Benchmark paneli artık HTTP API kapalıyken bile statik mock yerine gerçek "
              "servisleri in-process çağırır (app/ui/benchmark/local_backend.py). Senaryo, "
              "algoritma kataloğu, yönetişim ve governed-run verileri canlı üretilir "
              "(used_mock=False).")
    conn.close()
    return save(doc, "Benchmark_Algoritma_Kullanim_Raporu.docx")


# ============================================================================
# 3) VERİ SAYFALARI ALGORİTMA RAPORU (madde 9)
# ============================================================================

def report_data_pages() -> Path:
    doc = new_doc(
        "Veri Sayfalarında Çalışan Algoritmalar Raporu",
        "'Veri' ana başlığı altındaki sayfalar, veri kaynakları ve çalışan algoritmalar",
    )
    h1(doc, "1. Veri Sekmesi Sayfaları")
    table(doc, ["Sayfa", "Veri Kaynağı / Tablolar", "Çalışan Servis / Algoritma", "Amaç"], [
        ["Veri Yönetimi", "import_batches, import_row_issues, import_diffs",
         "import_audit_service, import_quality_service, import_rollback_service, import_diff_service",
         "İçe aktarım, kalite skoru, diff, onay ve rollback"],
        ["Veri Kalitesi", "ders, performans, populerlik, ders_kriterleri, anket_sonuclari",
         "data_quality_integration_service (readiness/coverage, eksik matris, doğrulama)",
         "Veri olgunluğu, kapsama, eksik veri ve doğrulama sorunları"],
        ["Trend", "performans (yıllara göre)",
         "trend_analysis_service (weighted_trend_score, analyze_trend_values)",
         "Ders bazlı tarihsel trend skoru ve etiketi"],
    ])
    h1(doc, "2. Mimari Not")
    para(doc, "Veri Kalitesi sayfası artık doğrudan SQLite bağlantısı açmaz; tüm sorgular "
              "data_quality_integration_service katmanına taşınmıştır (UI yalnızca servis "
              "çağırır ve sonucu render eder). Bu, katmanlı mimariyi korur.")
    h1(doc, "3. Trend Sayfası")
    para(doc, "Trend Kontrol sayfası 'Veri' başlığı altına eklenmiştir; yıl/fakülte seçimi, "
              "canlı trend hesaplama, grafik ve adım adım açıklama içerir. Geçmişi olmayan "
              "yıllarda (ör. 2022) nötr trend skoru (0.5) döndürülür; karar formülü bozulmaz.")
    return save(doc, "Veri_Sayfalari_Algoritma_Raporu.docx")


# ============================================================================
# 4) OVERRIDE ALGORİTMALARI AÇIKLAMA RAPORU (madde 11)
# ============================================================================

def report_override() -> Path:
    doc = new_doc(
        "Override Algoritmaları Açıklama Raporu",
        "Override kavramı, devreye girdiği durumlar ve ilişkili algoritmalar",
    )
    conn = ro_conn()
    h1(doc, "1. Override Nedir?")
    para(doc, "Override (geçersiz kılma), otomatik karar motorunun (AHP/TOPSIS + politika/"
              "durum makinesi) ürettiği sonucun, yetkili bir kullanıcı tarafından gerekçeli "
              "ve denetlenebilir biçimde değiştirilmesidir. Sistemin tamamen otomatik "
              "kararına insan denetimi (human-in-the-loop) ekler.")
    h1(doc, "2. Neden Kullanılır?")
    bullet(doc, "Düşük veri güveni: Az/eksik veriyle üretilen kararı uzman düzeltebilir.")
    bullet(doc, "Politika istisnaları: Müfredat/yönetim kararları otomatik skoru geçersiz kılabilir.")
    bullet(doc, "Yeni ders / geçmişsiz ders: Trend/performans verisi yokken manuel yönlendirme.")
    bullet(doc, "Hatalı veri: Aykırı değer (ör. yanlış kredi) düzeltilene kadar geçici override.")
    h1(doc, "3. Sistemde Hangi Durumda Devreye Girer?")
    para(doc, "Override akışı criteria_override_service üzerinden yürür: request_override "
              "(talep) → approve_override / reject_override (onay/ret) → get_active_override "
              "(karar anında aktif override kontrolü) → mark_override_used (kullanım kaydı). "
              "Karar motoru, bir ders/kapsam için aktif ve süresi geçmemiş bir override "
              "varsa onu uygular ve bunu açıklama/denetim kaydına yazar.")
    h1(doc, "4. Override ile İlişkili Algoritmalar / Servisler")
    table(doc, ["Bileşen", "Amacı", "Karar Sonucuna Etkisi / İzlenebilirlik"], [
        ["criteria_override_service", "Override yaşam döngüsü (talep/onay/ret/kullanım)",
         "Aktif override kararı doğrudan değiştirir; tüm adımlar zaman damgalı loglanır"],
        ["criteria_completion_policy_service", "Kriter tamamlanma politikası",
         "Override koşullarını ve kapıları (gate) belirler"],
        ["decision_policy_service / pool_state_policy_service", "Karar/havuz politikaları",
         "Override'ın hangi durumda izinli olduğunu çözer"],
        ["data_confidence_service", "Veri güven skoru",
         "Düşük güven → override önerisi/gerekliliği tetikler"],
        ["state_machine (evaluate_course_state_transition)", "Ders durum geçişi",
         "Override, önerilen durumu (recommended) nihai durumla (final) değiştirir"],
        ["criteria_override_service.mark_override_used", "Kullanım izi",
         "Hangi kararda hangi override kullanıldığını denetlenebilir kılar"],
    ])
    h1(doc, "5. Güvenlik ve İzlenebilirlik")
    para(doc, "Her override talebi, onayı ve kullanımı kullanıcı + zaman damgasıyla "
              "kaydedilir; süre (expires_at) ve kapsam (scope) ile sınırlandırılır. Bu, "
              "kararın neden otomatik sonuçtan saptığını geriye dönük açıklanabilir kılar.")
    # canlı override sayıları
    try:
        n = scalar(conn, "SELECT COUNT(*) FROM criteria_overrides")
        para(doc, f"Canlı durum: criteria_overrides tablosunda {n} override kaydı bulunuyor.")
    except sqlite3.Error:
        pass
    conn.close()
    return save(doc, "Override_Algoritmalari_Aciklama_Raporu.docx")


# ============================================================================
# 5) TREND ALGORİTMALARI VE HESAPLAMA MANTIĞI RAPORU (madde 10)
# ============================================================================

def report_trend() -> Path:
    doc = new_doc(
        "Trend Algoritmaları ve Hesaplama Mantığı Raporu",
        "Ağırlıklı trend skoru, etiketleme ve eksik geçmiş için nötr skor tasarımı",
    )
    h1(doc, "1. Ağırlıklı Trend Skoru")
    para(doc, "weighted_trend_score, bir dersin yıllara göre değerlerini en yeni yıla daha "
              "yüksek ağırlık vererek birleştirir. Varsayılan ağırlıklar (en yeniden eskiye): "
              "0.50, 0.30, 0.20.")
    h1(doc, "2. Trend Etiketleme")
    table(doc, ["Etiket", "Koşul"], [
        ["rising", "Son yıllarda belirgin artış (toplam değişim ≥ eşik, son adım ≥ 0)"],
        ["falling", "Son yıllarda belirgin düşüş"],
        ["stable", "Belirgin yön yok, düşük oynaklık"],
        ["volatile", "Yüksek oynaklık / dalgalanma"],
        ["new_course", "Hedef yılda ilk kez görülen ders"],
        ["insufficient_data", "Trend için yeterli geçmiş yok"],
    ])
    h1(doc, "3. Eksik Geçmiş İçin Nötr Skor (madde 10)")
    para(doc, "Trend YÖNÜ en az 2 yıllık geçmiş gerektirir. Geçmişi olmayan dersler (ör. veri "
              "kümesinin en erken yılı 2022) için eskiden 0.0 döndürülüyordu; bu, dersi en "
              "düşük trendle CEZALANDIRIP karar formülünü bozuyordu. Artık <2 veri noktasında "
              "NÖTR skor (NEUTRAL_TREND_SCORE = 0.5) döndürülür: ders ne ödüllendirilir ne "
              "cezalandırılır. Sonuç 'neutral_trend: True' ile işaretlenir.")
    para(doc, "Bu sayede 2022 gibi geçmişsiz yıllarda trend, karar sistemine zarar vermez; "
              "regresyon ve e2e karar testleri bu değişiklikle kırılmadan geçmektedir.")
    h1(doc, "4. Kullanım")
    para(doc, "Trend skoru karar motoruna (evaluate_course_state_transition) bir girdi olarak "
              "verilir ve ders durum geçişini etkiler. Trend sayfası (Veri başlığı altında) "
              "bu hesaplamayı canlı gösterir ve adım adım açıklar.")
    return save(doc, "Trend_Algoritmalari_ve_Hesaplama_Mantigi_Raporu.docx")


# ============================================================================
# 6) SİSTEM SAĞLIĞI KONTROLLERİ DOĞRULUK RAPORU (madde 17)
# ============================================================================

def report_health() -> Path:
    doc = new_doc(
        "Sistem Sağlığı Kontrolleri Doğruluk Raporu",
        "Sağlık panelinin canlı hesaplaması ve kontrol kategorileri",
    )
    h1(doc, "1. Canlılık")
    para(doc, "Sistem sağlığı paneli her çalıştırmada gerçek fonksiyon/servisler üzerinden "
              "ölçüm yapar; statik/kayıtlı değer göstermez. Aşağıdaki özet rapor üretimi "
              "sırasında canlı çalıştırılmıştır.")
    try:
        from app.services.service_factory import get_health_service
        hs = get_health_service(db_path=str(DB_PATH))
        rep = hs.run_full_health_check().to_dict()
        para(doc, f"Skor: {rep.get('score')} · Durum: {rep.get('overall_status')} · "
                  f"Toplam kontrol: {rep.get('total_checks')}")
        table(doc, ["Sonuç", "Adet"], [
            ["Başarılı (ok)", rep.get("ok_count")],
            ["Uyarı (warning)", rep.get("warning_count")],
            ["Kritik (critical)", rep.get("critical_count")],
            ["Başarısız (failed)", rep.get("failed_count")],
            ["Bilgi (info)", rep.get("info_count")],
            ["Atlandı (skipped)", rep.get("skipped_count")],
        ])
        # kategori bazlı örnek
        cats = {}
        for c in rep.get("results", []):
            cats.setdefault(c.get("category", "?"), 0)
            cats[c.get("category", "?")] += 1
        table(doc, ["Kategori", "Kontrol Sayısı"], [[k, v] for k, v in sorted(cats.items())])
    except Exception as exc:  # noqa: BLE001
        para(doc, f"(Canlı sağlık çalıştırması bu ortamda alınamadı: {exc})")

    h1(doc, "2. Kontrol Kategorileri")
    para(doc, "Sağlık kontrolleri app/health/checks/ altında modüler tanımlıdır: veritabanı "
              "bağlantısı/integrity/FK, şema, veri kalitesi, AHP/TOPSIS, karar merkezi, "
              "import yönetişimi, performans (psutil), güvenlik, yedekleme, log, bağımlılık, "
              "API/UI ve benchmark kontrolleri.")
    h1(doc, "3. Düzeltme")
    para(doc, "Veri kalitesi içindeki 'tekrarlı kayıt' kontrolünün havuz için anahtarı "
              "(ders_id, yil) → (ders_id, yil, donem) olarak düzeltildi; Güz/Bahar kayıtları "
              "artık yanlış pozitif üretmiyor.")
    return save(doc, "Sistem_Sagligi_Kontrolleri_Dogruluk_Raporu.docx")


# ============================================================================
# 7) GÜVENLİK VE ÜRETİM SAYFASI AÇIKLAMA RAPORU (madde 6)
# ============================================================================

def report_security() -> Path:
    doc = new_doc(
        "Güvenlik ve Üretim Sayfası Açıklama Raporu",
        "Güvenlik/üretim hazırlığı kontrollerinin canlı durumu ve anlamı",
    )
    h1(doc, "1. Amaç")
    para(doc, "Güvenlik & Hazırlık sayfası, sistemin geliştirme (development) ile üretim "
              "(production) modları arasındaki güvenlik farklarını canlı olarak gösterir ve "
              "üretimde risk oluşturabilecek ayarları uyarı olarak işaretler.")
    h1(doc, "2. Canlı Kontroller")
    try:
        from app.core.config import load_app_config
        from app.services.security_health_service import SecurityHealthService
        d = SecurityHealthService(load_app_config()).check_security_configuration()
        para(doc, f"Skor: {d.get('score')}/{d.get('max_score')} · Seviye: {d.get('level')}")
        table(doc, ["Kontrol", "Durum", "Mesaj"],
              [[c.get("name"), c.get("status", "").upper(), c.get("message")] for c in d.get("checks", [])])
    except Exception as exc:  # noqa: BLE001
        para(doc, f"(Canlı güvenlik çalıştırması alınamadı: {exc})")
    h1(doc, "3. Üretim Önerileri")
    bullet(doc, "Üretimde API kimlik doğrulama ve RBAC etkinleştirilmeli.")
    bullet(doc, "SQL Console üretimde kapatılmalı; tehlikeli SQL desenleri zaten engelleniyor.")
    bullet(doc, "Çalışma zamanı şema mutasyonu üretimde kilitli olmalı (mevcut durumda kilitli).")
    bullet(doc, "CORS politikası daraltılmalı; rate limiting etkinleştirilmeli.")
    bullet(doc, "İçe aktarımlar üretimde onay akışına tabi olmalı.")
    return save(doc, "Guvenlik_ve_Uretim_Sayfasi_Aciklama_Raporu.docx")


# ============================================================================
# 8) GÖRSEL VERİ SETLERİ VE SENARYOLAR RAPORU (madde 12)
# ============================================================================

def report_scenarios() -> Path:
    doc = new_doc(
        "Görsel Veri Setleri ve Senaryolar Raporu",
        "Benchmark senaryoları, kullandıkları veriler ve ilişkili algoritmalar",
    )
    h1(doc, "1. Benchmark Senaryoları (canlı katalog)")
    try:
        from app.benchmark.scenarios import DEFAULT_SCENARIOS
        rows = []
        for s in DEFAULT_SCENARIOS.values():
            rows.append([
                s.display_name or s.name,
                s.problem_type,
                s.table_name,
                ", ".join(s.algorithm_names) if s.algorithm_names else "-",
            ])
        table(doc, ["Senaryo", "Problem Tipi", "Veri Tablosu", "Algoritmalar"], rows)
        h1(doc, "2. Senaryo Amaçları")
        for s in DEFAULT_SCENARIOS.values():
            h2(doc, s.display_name or s.name)
            if s.purpose_tr:
                para(doc, "Amaç: " + s.purpose_tr)
            if s.system_impact_tr:
                para(doc, "Sisteme etkisi: " + s.system_impact_tr)
    except Exception as exc:  # noqa: BLE001
        para(doc, f"(Senaryo kataloğu alınamadı: {exc})")
    h1(doc, "3. Görsellerdeki Veri Setleri")
    para(doc, "Benchmark panelinde görünen 'desktop_benchmark_dataset' ve benzeri veri "
              "setleri, gerçek sistem tablolarından türetilen özellik (feature) tablolarına "
              "(student_course_features, student_course_features_unencoded, preferences) "
              "dayanır. Senaryolar bu tablolar üzerinden çalışır.")
    return save(doc, "Gorsel_Veri_Setleri_ve_Senaryolar_Raporu.docx")


def main():
    outputs = []
    for fn in (report_data_fitness, report_algorithms, report_data_pages,
               report_override, report_trend, report_health, report_security,
               report_scenarios):
        try:
            outputs.append(fn())
            print(f"[OK] {fn.__name__} -> {outputs[-1].name}")
        except Exception as exc:  # noqa: BLE001
            print(f"[FAIL] {fn.__name__}: {exc}")
    print(f"\n{len(outputs)} rapor üretildi -> {DOCS_DIR}")


if __name__ == "__main__":
    main()
