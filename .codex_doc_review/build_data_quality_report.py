from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\dadil\Documents\Yazilimsal_Projeler\BitirmeProje\Adil_Secmeli_Python")
OUT = ROOT / "reports" / "VERI_KALITE_SKORU_MATEMATIKSEL_ANALIZ_2026-06-18.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF4CC"
PALE_RED = "FDECEA"
GREEN = "E6F4EA"


def set_font(run, name="Calibri", size=11, bold=False, color="000000", italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, bold=False, color="000000", align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(str(text))
    set_font(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, header_fill=LIGHT, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        shade(table.rows[0].cells[i], header_fill)
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
    set_repeat_table_header(table.rows[0])
    for ridx, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if ridx % 2 == 1:
                shade(cells[i], "FAFBFC")
            align = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_lead):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_callout(doc, title, text, fill=PALE_BLUE, title_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, bold=True, color=title_color, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, size=10.2)
    set_table_geometry(table, [9360])
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in (
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("ADİL SEÇMELİ • TEKNİK İNCELEME")
    set_font(hr, size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    tr = title.add_run("Veri Kalite Skoru Matematiksel Analiz Raporu")
    set_font(tr, size=24, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sr = subtitle.add_run("%69 sonucunun kaynağı, kullanılan algoritmalar ve 71 satırlık importun yeniden hesabı")
    set_font(sr, size=12, color=MUTED)

    add_table(
        doc,
        ["Alan", "Değer"],
        [
            ("İncelenen kayıt", "Import batch #6"),
            ("Dosya", "2022_anket_tercih_veri_setii.xlsx"),
            ("Import türü", "Anket (survey)"),
            ("Kapsam", "2022 • Fakülte ID 5 • 16 işlenen satır"),
            ("Rapor tarihi", "18 Haziran 2026"),
            ("Kaynak", "data/adil_secmeli.db ve uygulama servis kodu"),
        ],
        [2400, 6960],
    )

    add_callout(
        doc,
        "Sonuç",
        "Arayüzde görülen %69 skoru matematiksel olarak doğrudur; ancak kullanılan 71 paydası bu fakülte kapsamı için doğru değildir. "
        "Batch #6 içinde 16 satırın tamamı başarıyla eşleşmiştir. Buna rağmen servis, fakülteye ait 16 satırı dosyanın tüm fakültelerdeki 71 satırına bölmüştür. "
        "Kapsam paydası 16 olsaydı mevcut formül aynı veriye %100 kalite skoru verirdi.",
        fill=PALE_GOLD,
        title_color="7A5A00",
    )

    add_heading(doc, "1. İncelenen sonucun kimliği", 1)
    add_body(
        doc,
        "Kullanıcı ekranındaki değerler veritabanındaki import_quality_checks kaydıyla birebir eşleşmektedir: "
        "quality_score=0,6901; successful_row_ratio=0,2254; matched_course_ratio=0,2254; "
        "valid_numeric_ratio=1,0; row_count=71; actual_row_count=16."
    )
    add_body(
        doc,
        "Önemli ayrım: Bu rapor müfredat veya kriter batch’ine değil, anket importuna aittir. Aynı veritabanında müfredat importu batch #1 olarak "
        "2022_Mufredat.xlsx adıyla kayıtlıdır ve kalite skoru 0,60’tır. Ekrandaki %69 ve 71 satır bilgisi ise batch #6 anket kaydını tanımlar."
    )

    add_heading(doc, "2. Kalite skoru hangi algoritmayla hesaplanıyor?", 1)
    add_body(
        doc,
        "Bu skor AHP, TOPSIS, ELECTRE, karar ağacı, rastgele orman veya başka bir makine öğrenmesi modeliyle hesaplanmaz. "
        "Kullanılan yöntem deterministik, ağırlıklı doğrusal puanlama formülüdür. Aynı girdiler her çalıştırmada aynı sonucu üretir."
    )
    add_callout(
        doc,
        "Genel formül",
        "Q = 0,20R + 0,20S + 0,20M + 0,15V + 0,10(1−D) + 0,10K + 0,05C",
        fill=PALE_BLUE,
    )
    doc.add_page_break()
    add_table(
        doc,
        ["Sembol", "Bileşen", "Tanım", "Ağırlık"],
        [
            ("R", "Zorunlu kolon", "Zorunlu başlık/alan hatası yoksa 1, aksi halde 0", "%20"),
            ("S", "Başarılı satır", "Başarılı kabul edilen satır / toplam satır", "%20"),
            ("M", "Ders eşleşmesi", "matched_ders_id dolu satır / toplam satır", "%20"),
            ("V", "Sayısal geçerlilik", "1 − (geçersiz + aralık dışı) / toplam", "%15"),
            ("D", "Tekrar cezası", "Tekrarlı satır / toplam; formülde 1−D kullanılır", "%10"),
            ("K", "Kapsam tutarlılığı", "invalid_scope yoksa 1, varsa 0", "%10"),
            ("C", "Tamlık", "1 − (eksik zorunlu + eşleşmeyen) / toplam", "%5"),
        ],
        [850, 1900, 5410, 1200],
        font_size=8.7,
    )

    add_heading(doc, "3. Satır ve eşleşme algoritmaları", 1)
    add_heading(doc, "3.1 Başarılı satır sınıflandırması", 2)
    add_body(
        doc,
        "Satır durumu matched, applied, ok veya success ise satır kalite değeri 1,00’dır. skipped_override veya warning ise 0,65’tir. "
        "Durum bunlardan biri değil fakat matched_ders_id mevcutsa 0,75; aksi halde 0,00 verilir. Başarı oranında değeri en az 0,65 olan satırlar sayılır."
    )
    add_heading(doc, "3.2 Ders eşleştirme sırası", 2)
    add_table(
        doc,
        ["Sıra", "Yöntem", "İşlem"],
        [
            ("1", "Ders kodu", "Küçük harf, Türkçe karakter dönüşümü ve boşluk normalizasyonundan sonra tam eşleşme"),
            ("2", "Ders adı", "Aynı metin normalizasyonundan sonra tam ad eşleşmesi"),
            ("3", "Normalize ad", "Harf ve rakam dışı karakterler kaldırıldıktan sonra tam anahtar eşleşmesi"),
            ("4", "Belirsizlik çözümü", "Birden fazla aday varsa seçili yıl kapsamındaki tek aday tercih edilir; hâlâ çoksa hata"),
        ],
        [800, 1900, 6660],
    )
    add_body(
        doc,
        "Bu mekanizma bulanık benzerlik veya olasılıksal tahmin kullanmaz. Levenshtein, cosine similarity veya makine öğrenmesi yoktur; normalleştirilmiş kesin eşleşme uygulanır."
    )

    doc.add_page_break()
    add_heading(doc, "4. %69 sonucunun adım adım yeniden hesabı", 1)
    add_body(doc, "Servisin toplam satır tanımı aşağıdaki gibidir:")
    add_callout(doc, "Payda", "N = max(kayıtlı gerçek satır sayısı, batch üzerinde bildirilen satır sayısı, 1) = max(16, 71, 1) = 71")
    add_table(
        doc,
        ["Bileşen", "Ham değer", "Oran", "Ağırlıklı katkı"],
        [
            ("Zorunlu kolonlar", "Eksik=0", "R=1", "0,20 × 1 = 0,200000"),
            ("Başarılı satırlar", "16 başarılı / 71", "S=0,225352", "0,20 × S = 0,045070"),
            ("Ders eşleşmesi", "16 eşleşmiş / 71", "M=0,225352", "0,20 × M = 0,045070"),
            ("Sayısal geçerlilik", "Hata=0; aralık dışı=0", "V=1", "0,15 × 1 = 0,150000"),
            ("Tekrar cezası", "Tekrar=0", "1−D=1", "0,10 × 1 = 0,100000"),
            ("Kapsam tutarlılığı", "invalid_scope=0", "K=1", "0,10 × 1 = 0,100000"),
            ("Tamlık", "Eksik=0; eşleşmeyen=0", "C=1", "0,05 × 1 = 0,050000"),
        ],
        [2200, 2600, 1700, 2860],
        font_size=8.7,
    )
    add_callout(
        doc,
        "Toplam",
        "Q = 0,200000 + 0,045070 + 0,045070 + 0,150000 + 0,100000 + 0,100000 + 0,050000 "
        "= 0,6901408. Servis dört ondalığa yuvarlar: 0,6901. Arayüz yüzdeyi tam sayıya yuvarlar: %69.",
        fill=GREEN,
        title_color="1B5E20",
    )

    add_heading(doc, "5. Neden sorun sayısı sıfırken başarı ve eşleşme %23?", 1)
    add_body(
        doc,
        "Bu görünüm ilk bakışta çelişkilidir; ancak iki farklı evren kullanıldığı için oluşur. Sorun sayıları yalnız batch #6 içinde saklanan 16 fakülte satırını inceler. "
        "Bu 16 satırın tamamı matched durumundadır; dolayısıyla eşleşmeyen, hata ve uyarı sayıları sıfırdır. Oranların paydası ise batch oluşturulurken dosyanın tamamından alınan 71’dir."
    )
    add_table(
        doc,
        ["Fakülte batch’i", "İşlenen/eşleşen", "Kullanılan payda", "Kalite skoru"],
        [
            ("Batch #2 • Fakülte 4", "19", "71", "%70,70"),
            ("Batch #3 • Fakülte 2", "12", "71", "%66,76"),
            ("Batch #4 • Fakülte 3", "10", "71", "%65,63"),
            ("Batch #5 • Fakülte 1", "14", "71", "%67,89"),
            ("Batch #6 • Fakülte 5", "16", "71", "%69,01"),
            ("Toplam", "71", "—", "—"),
        ],
        [3200, 2100, 1900, 2160],
    )
    add_callout(
        doc,
        "Kök neden",
        "Anket importu önce create_import_batch çağrısında Excel metadata satır sayısını (71) batch’e yazar; daha sonra satırları seçili fakülteye göre filtreler. "
        "Kalite servisi total_rows = max(actual_row_count, declared_row_count, 1) kullandığı için fakülte batch’lerinin tamamında ortak 71 paydası kalır.",
        fill=PALE_RED,
        title_color="9B1C1C",
    )

    add_heading(doc, "6. Doğru kapsam paydasıyla karşılaştırma", 1)
    add_body(
        doc,
        "Batch #6’nın kalite değerlendirmesi fakülte kapsamındaki gerçek 16 satır üzerinden yapılırsa S=16/16=1 ve M=16/16=1 olur. Diğer tüm bileşenler zaten 1’dir."
    )
    add_callout(
        doc,
        "Karşı-olgusal hesap",
        "Qdoğru = 0,20 + 0,20×1 + 0,20×1 + 0,15×1 + 0,10×1 + 0,10×1 + 0,05×1 = 1,00 (%100, Çok iyi)",
        fill=GREEN,
        title_color="1B5E20",
    )
    add_body(
        doc,
        "Bu %100 sonucu, anket değerlerinin akademik olarak kusursuz olduğu anlamına gelmez. Yalnız mevcut kalite formülünün kontrol ettiği şema, eşleşme, sayısal geçerlilik, tekrar, kapsam ve zorunlu alan kurallarında hata bulunmadığını ifade eder."
    )

    add_heading(doc, "7. Kalite seviyeleri ve karar algoritmalarıyla ilişkisi", 1)
    add_table(
        doc,
        ["Skor aralığı", "Seviye", "Import davranışı"],
        [
            ("Q ≥ 0,80", "high • Çok iyi", "Doğrulanır; otomatik aktivasyona uygundur"),
            ("0,55 ≤ Q < 0,80", "medium • Kullanılabilir", "Doğrulanır; mevcut akışta otomatik aktivasyona uygundur"),
            ("Q < 0,55", "low • Riskli", "pending_review durumuna alınır; otomatik aktive edilmez"),
        ],
        [1900, 2500, 4960],
    )
    add_body(
        doc,
        "Kalite skoru import yönetişim kapısıdır. Daha sonra çalışan AHP kriter ağırlıklandırması, TOPSIS sıralaması, ELECTRE TRI-B sınıflandırması, trend analizi ve destekleyici makine öğrenmesi modelleri bu %69’u üretmez. "
        "Bu modeller yalnız kabul edilip etkinleştirilen kriter/anket verilerini kendi karar hesaplarında kullanabilir."
    )

    doc.add_page_break()
    add_heading(doc, "8. Teknik değerlendirme ve önerilen düzeltmeler", 1)
    add_table(
        doc,
        ["Öncelik", "Bulgu", "Önerilen değişiklik"],
        [
            ("P0", "Fakülte batch paydası global dosya satır sayısı", "Batch row_count değerini fakülte filtresinden sonra len(survey_rows) olarak yaz"),
            ("P1", "Ekran 71 satırı ‘incelenen’ diye gösteriyor", "row_count ve actual_row_count değerlerini ayrı ayrı göster"),
            ("P1", "55 kapsam dışı satır görünmez", "excluded_by_scope_count=55 alanını kalite özetine ekle; hata olarak sayma"),
            ("P2", "Başarı ve eşleşme aynı sinyali iki kez ağırlıklandırıyor", "Import türüne göre ağırlıkları gözden geçir veya iki metriğin anlamını ayrıştır"),
            ("P2", "Kalite %100 akademik doğruluk sanılabilir", "Arayüzde ‘teknik import kalitesi’ ifadesini kullan"),
            ("P0", "Regresyon riski", "Çok fakülteli 71 satır / fakülte bazlı batch testi ekle; beklenen batch #6 paydası 16 olsun"),
        ],
        [950, 3300, 5110],
        font_size=8.5,
    )

    add_heading(doc, "9. Kaynak kod ve veri kanıtları", 1)
    add_table(
        doc,
        ["Kaynak", "İlgili bölüm"],
        [
            ("app/services/import_quality_service.py", "Seviye eşikleri: 50–55; oranlar: 115–120; ağırlıklı formül: 122–132"),
            ("app/services/survey_import_service.py", "Global metadata row_count: yaklaşık 1038–1047; fakülte filtresi: yaklaşık 1071"),
            ("app/services/course_matcher.py", "Metin normalizasyonu: 30–47; kesin eşleştirme sırası: 151–204"),
            ("app/ui/tabs/data_management_page.py", "Yüzde ve kullanıcı metni biçimlendirmesi: yaklaşık 1147–1193"),
            ("data/adil_secmeli.db", "import_batches #1–#6, import_quality_checks #12, survey_import_rows batch #6"),
        ],
        [3600, 5760],
        font_size=8.8,
    )

    add_heading(doc, "10. Nihai hüküm", 1)
    add_body(
        doc,
        "Gösterilen %69, mevcut kodun uyguladığı formüle göre hatasız hesaplanmıştır; fakat veri kapsamını temsil etmeyen bir payda nedeniyle yanıltıcı derecede düşüktür. "
        "Batch #6 için teknik gerçeklik şudur: 16/16 satır eşleşmiş, sayısal ve yapısal hata bulunmamış, fakat oran hesabında 16 yerine tüm dosyanın 71 satırı kullanılmıştır. "
        "Bu nedenle sorun verinin kalitesinden çok çok-fakülteli anket importunun batch kapsamı ile kalite paydasının uyumsuzluğudur."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
